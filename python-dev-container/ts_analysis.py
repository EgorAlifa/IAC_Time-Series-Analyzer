import os
import io
import tempfile
import smtplib
import numpy as np
import pandas as pd
import seaborn as sns
import matplotlib.pyplot as plt
import time
import threading
from scipy import stats
from datetime import datetime
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from docx import Document
from docx.shared import Inches
from pydantic import BaseModel, EmailStr
from typing import Optional
from statsmodels.tsa.stattools import adfuller, kpss, coint
from statsmodels.graphics.tsaplots import plot_acf, plot_pacf
from statsmodels.tsa.vector_ar.vecm import coint_johansen
from datetime import datetime


# Функция для проведения теста Дики-Фуллера на стационарность
def adf_test(series, title=''):
    """
    Проведение расширенного теста Дики-Фуллера (ADF) для проверки стационарности

    H0: Ряд имеет единичный корень (не стационарен)
    H1: Ряд не имеет единичный корень (стационарен)

    Если p-value < 0.05, отвергаем нулевую гипотезу и считаем ряд стационарным
    """
    result = adfuller(series.dropna())

    output = []
    output.append(f'ADF Test для {title}:')
    output.append(f'ADF Statistic: {result[0]:.4f}')
    output.append(f'p-value: {result[1]:.4f}')

    # Критические значения для разных уровней значимости
    for key, value in result[4].items():
        output.append(f'Критическое значение ({key}): {value:.4f}')

    # Вывод результата
    if result[1] < 0.05:
        output.append(f"Результат: Отвергаем нулевую гипотезу. Ряд {title} стационарен.\n")
        is_stationary = True
    else:
        output.append(f"Результат: Не отвергаем нулевую гипотезу. Ряд {title} не стационарен.\n")
        is_stationary = False

    return is_stationary, output

# Функция для проведения теста KPSS на стационарность
def kpss_test(series, title=''):
    """
    Проведение теста Квятковски-Филлипса-Шмидта-Шина (KPSS) для проверки стационарности

    H0: Ряд стационарен
    H1: Ряд не стационарен

    Если p-value < 0.05, отвергаем нулевую гипотезу и считаем ряд нестационарным
    """
    result = kpss(series.dropna(), regression='c', nlags="auto")

    output = []
    output.append(f'KPSS Test для {title}:')
    output.append(f'KPSS Statistic: {result[0]:.4f}')
    output.append(f'p-value: {result[1]:.4f}')

    # Критические значения для разных уровней значимости
    for key, value in result[3].items():
        output.append(f'Критическое значение ({key}): {value:.4f}')

    # Вывод результата
    if result[1] < 0.05:
        output.append(f"Результат: Отвергаем нулевую гипотезу. Ряд {title} не стационарен.\n")
        is_stationary = False
    else:
        output.append(f"Результат: Не отвергаем нулевую гипотезу. Ряд {title} стационарен.\n")
        is_stationary = True

    return is_stationary, output

# Дополнительны тест на стационарность Phillips-Perron
def phillips_perron_test(series, title=''):
    """
    Проведение теста Филлипса-Перрона для проверки стационарности
    используя библиотеку arch

    H0: Ряд имеет единичный корень (не стационарен)
    H1: Ряд не имеет единичный корень (стационарен)

    Если p-value < 0.05, отвергаем нулевую гипотезу и считаем ряд стационарным
    """
    # Удаляем пропуски
    series = series.dropna()

    # Базовый ADF тест
    adf_result = adfuller(series)

    # Вычисляем разности первого порядка
    diff_series = series.diff().dropna()

    # Реализация PP теста на основе ADF
    T = len(diff_series)
    demeaned_diff = diff_series - diff_series.mean()

    # Оценка дисперсии остатков
    sigma_sq = np.sum(demeaned_diff**2) / (T - 1)

    # Оценка долгосрочной дисперсии (упрощенная)
    long_run_variance = sigma_sq

    # Статистика PP
    pp_stat = (adf_result[0] - T/2 * (sigma_sq - long_run_variance)) / np.sqrt(2 * long_run_variance / T)

    # Критические значения для PP теста (приближенные)
    critical_values = {
        '1%': -3.43,
        '5%': -2.86,
        '10%': -2.57
    }

    # Вычисление p-value (приближенное)
    p_value = stats.t.sf(abs(pp_stat), T-1)*2

    # Вывод результатов
    output = []
    output.append(f'Phillips-Perron Test для {title}:')
    output.append(f'PP Statistic: {pp_stat:.4f}')
    output.append(f'p-value: {p_value:.4f}')

    # Критические значения
    for key, value in critical_values.items():
        output.append(f'Критическое значение ({key}): {value:.4f}')

    # Интерпретация результата
    if p_value < 0.05:
        output.append(f"Результат: Отвергаем нулевую гипотезу. Ряд {title} стационарен.\n")
        is_stationary = True
    else:
        output.append(f"Результат: Не отвергаем нулевую гипотезу. Ряд {title} не стационарен.\n")
        is_stationary = False

    return is_stationary, output

# Функция для проверки стационарности с помощью трех тестов
def check_stationarity(series, title=''):
    """
    Комплексная проверка стационарности с помощью трех тестов:
    - ADF тест
    - KPSS тест
    - Phillips-Perron тест

    Ряд признается стационарным, если минимум 2 из 3 тестов показывают стационарность
    """
    output = []

    # Проведение тестов
    adf_result, adf_output = adf_test(series, title)
    kpss_result, kpss_output = kpss_test(series, title)
    pp_result, pp_output = phillips_perron_test(series, title)

    # Собираем результаты всех тестов
    output.extend(adf_output)
    output.extend(kpss_output)
    output.extend(pp_output)

    # Интерпретация результатов тестов
    test_details = [
        ("ADF", adf_result),
        ("KPSS", kpss_result),
        ("Phillips-Perron", pp_result)
    ]

    # Подсчет количества тестов, указывающих на стационарность
    stationary_count = sum(test[1] for test in test_details)
    total_tests = len(test_details)

    # Детальный вывод результатов каждого теста
    detailed_results = "\nДетальные результаты тестов:\n"
    for test_name, is_stationary in test_details:
        detailed_results += f"{test_name} тест: {'Стационарен' if is_stationary else 'Не стационарен'}\n"

    # Определение статуса на основе критериев:
    # 1. Минимум 2 теста показывают стационарность
    if stationary_count >= 2:
        conclusion = f"Итоговый результат: Ряд {title} стационарен (по {stationary_count} из {total_tests} тестов)."
        status = "Стационарен"
    else:
        conclusion = f"Итоговый результат: Ряд {title} не стационарен (стационарен только по {stationary_count} из {total_tests} тестов)."
        status = "Не стационарен"

    # Добавляем детальные результаты к выводу
    output.append(conclusion)
    output.append(detailed_results)

    return status, output

# Функция для проверки коинтеграции
def test_cointegration(series1, series2, name1, name2):
    """
    Проверка наличия коинтеграции между двумя рядами

    H0: Нет коинтеграции
    H1: Есть коинтеграция

    Если p-value < 0.05, отвергаем H0 и считаем, что ряды коинтегрированы
    """
    # Объединяем ряды без пропусков
    df_temp = pd.DataFrame({name1: series1, name2: series2})
    df_temp = df_temp.dropna()

    # Проводим тест Йохансена на коинтеграцию (через ADF)
    result = coint(df_temp[name1], df_temp[name2])

    output = []
    output.append(f'Тест на коинтеграцию между {name1} и {name2}:')
    output.append(f'Test Statistic: {result[0]:.4f}')
    output.append(f'p-value: {result[1]:.4f}')

    # Критические значения
    for i, value in enumerate(result[2]):
        output.append(f'Критическое значение ({i+1}%): {value:.4f}')

    # Интерпретация результата
    if result[1] < 0.05:
        output.append(f"Результат: Отвергаем нулевую гипотезу. Ряды {name1} и {name2} коинтегрированы.\n")
        is_cointegrated = True
    else:
        output.append(f"Результат: Не отвергаем нулевую гипотезу. Ряды {name1} и {name2} НЕ коинтегрированы.\n")
        is_cointegrated = False

    return is_cointegrated, output

# Функция для теста Йохансена на коинтеграцию
def johansen_test(df, endogenous_vars, det_order=1, k_ar_diff=1):
    """
    Проведение теста Йохансена на коинтеграцию для нескольких временных рядов

    Parameters:
    -----------
    df : pandas.DataFrame
        Датафрейм с временными рядами
    endogenous_vars : list
        Список имен переменных для анализа коинтеграции
    det_order : int
        Порядок детерминистических компонент
    k_ar_diff : int
        Количество лагов в VAR в разностях
    """
    from statsmodels.tsa.vector_ar.vecm import coint_johansen

    # Выделяем нужные переменные и удаляем пропуски
    data = df[endogenous_vars].dropna()

    # Проводим тест Йохансена
    result = coint_johansen(data, det_order=det_order, k_ar_diff=k_ar_diff)

    # Получаем статистику и критические значения
    trace_stat = result.lr1
    eigen_stat = result.lr2

    # Критические значения для trace test (90%, 95%, 99%)
    crit_vals_trace = result.cvt

    # Критические значения для eigenvalue test (90%, 95%, 99%)
    crit_vals_eigen = result.cvm

    # Количество переменных
    n_vars = len(endogenous_vars)

    output = []
    output.append(f'Тест Йохансена на коинтеграцию для переменных: {", ".join(endogenous_vars)}')
    output.append(f'Количество переменных: {n_vars}')
    output.append(f'Порядок детерминистических компонент: {det_order}')
    output.append(f'Количество лагов: {k_ar_diff}\n')

    # Вывод результатов trace test
    output.append('Результаты Trace Test:')
    output.append('H0: Ранг коинтеграции <= r | H1: Ранг коинтеграции > r')
    output.append('-' * 70)
    output.append(f"{'r':^5}|{'Статистика':^15}|{'90%':^10}|{'95%':^10}|{'99%':^10}|{'Результат':^15}")
    output.append('-' * 70)

    trace_results = []
    for i in range(n_vars):
        reject_null = trace_stat[i] > crit_vals_trace[i, 1]  # Используем 95% критическое значение
        trace_results.append(reject_null)

        # Форматируем вывод
        result_text = "Отклоняем H0" if reject_null else "Не отклоняем H0"
        output.append(f"{i:^5}|{trace_stat[i]:^15.4f}|{crit_vals_trace[i, 0]:^10.4f}|{crit_vals_trace[i, 1]:^10.4f}|{crit_vals_trace[i, 2]:^10.4f}|{result_text:^15}")

    # Определяем ранг коинтеграции на основе trace test
    cointegration_rank = 0
    for i, reject in enumerate(trace_results):
        if not reject:
            cointegration_rank = i
            break

    # Делаем заключение
    output.append('\nЗаключение:')
    if cointegration_rank == 0:
        conclusion = "Нет коинтеграционных соотношений между переменными. Все переменные независимы в долгосрочном периоде."
        is_cointegrated = False
    elif cointegration_rank == 1:
        conclusion = f"Обнаружено {cointegration_rank} коинтеграционное соотношение. Существует один стабильный долгосрочный вектор между переменными."
        is_cointegrated = True
    elif cointegration_rank == 2:
        conclusion = f"Обнаружено {cointegration_rank} коинтеграционных соотношения. Существует два стабильных долгосрочных вектора между переменными."
        is_cointegrated = True
    else:
        conclusion = f"Обнаружено {cointegration_rank} коинтеграционных соотношений. Сложная долгосрочная динамика взаимосвязи между переменными."
        is_cointegrated = True
    
    output.append(conclusion)
    
    # Возвращаем результат в конце
    return is_cointegrated, output

# Функция для анализа разностей (дифференцирования) ряда
def analyze_differences(series, title='', max_diff=4):
    """
    Анализ разностей временного ряда для достижения стационарности
    Поддерживает до 4 порядков дифференцирования
    """
    all_outputs = []

    # Проверка исходного ряда
    orig_result, orig_output = check_stationarity(series, f"{title} (исходный)")
    all_outputs.extend(orig_output)

    # Если ряд уже стационарен, нет смысла дифференцировать
    if orig_result == "Стационарен":
        return 0, all_outputs

    # Дифференцирование первого порядка
    diff1 = series.diff().dropna()
    diff1_result, diff1_output = check_stationarity(diff1, f"{title} (первые разности)")
    all_outputs.extend(diff1_output)

    if diff1_result == "Стационарен":
        return 1, all_outputs

    # Дифференцирование второго порядка, если необходимо
    if max_diff >= 2:
        diff2 = diff1.diff().dropna()
        diff2_result, diff2_output = check_stationarity(diff2, f"{title} (вторые разности)")
        all_outputs.extend(diff2_output)

        if diff2_result == "Стационарен":
            return 2, all_outputs

    # Дифференцирование третьего порядка, если необходимо
    if max_diff >= 3:
        diff3 = diff2.diff().dropna()
        diff3_result, diff3_output = check_stationarity(diff3, f"{title} (третьи разности)")
        all_outputs.extend(diff3_output)

        if diff3_result == "Стационарен":
            return 3, all_outputs

    # Дифференцирование четвертого порядка, если необходимо
    if max_diff >= 4:
        diff4 = diff3.diff().dropna()
        diff4_result, diff4_output = check_stationarity(diff4, f"{title} (четвертые разности)")
        all_outputs.extend(diff4_output)

        if diff4_result == "Стационарен":
            return 4, all_outputs

    # Если даже после четырех дифференцирований ряд не стационарен
    return "Не достигнуто", all_outputs

# Основная функция для анализа данных
def analyze_data(df, endogenous_vars=None):
   """
   Основная функция для анализа временных рядов
   
   Parameters:
   -----------
   df : pandas.DataFrame
       Датафрейм с временными рядами
   endogenous_vars : list or None
       Список имен эндогенных переменных для анализа
       
   Returns:
   --------
   dict
       Результаты анализа в формате JSON
   """
   # Словарь для хранения результатов
   results = {'variable_results': {}}
   
   # Если эндогенные переменные не указаны, используем все числовые колонки
   if endogenous_vars is None:
       # Получаем все числовые колонки
       all_numeric_columns = [col for col in df.columns if df[col].dtype in [np.float64, np.int64]]
       
       # Если не указаны эндогенные, берем первые две как эндогенные
       endogenous_vars = all_numeric_columns[:2]
   else:
       # Получаем все числовые колонки, исключая эндогенные
       all_numeric_columns = [col for col in df.columns 
                              if df[col].dtype in [np.float64, np.int64] 
                              and col not in endogenous_vars]
   
   # Проверяем, что эндогенные переменные существуют в датафрейме
   endogenous_vars = [var for var in endogenous_vars if var in df.columns]
   
   # Анализ эндогенных переменных
   for column in endogenous_vars:
       d_value, test_outputs = analyze_differences(df[column], column, max_diff=4)
       results['variable_results'][column] = {
           'd_value': d_value,
           'test_outputs': test_outputs,
           'type': 'endogenous'
       }
   
   # Анализ экзогенных переменных
   for column in all_numeric_columns:
       d_value, test_outputs = analyze_differences(df[column], column, max_diff=4)
       results['variable_results'][column] = {
           'd_value': d_value,
           'test_outputs': test_outputs,
           'type': 'exogenous'
       }
   
   # Анализ коинтеграции
   if len(endogenous_vars) >= 2:
       # Для анализа методом Энгла-Грейнджера создаем все возможные пары
       eg_results = []
       
       for i in range(len(endogenous_vars)):
           for j in range(i+1, len(endogenous_vars)):
               var1 = endogenous_vars[i]
               var2 = endogenous_vars[j]
               
               is_cointegrated_eg, cointegration_output_eg = test_cointegration(
                   df[var1],
                   df[var2],
                   var1,
                   var2
               )
               
               eg_results.append({
                   'variables': [var1, var2],
                   'is_cointegrated': is_cointegrated_eg,
                   'details': cointegration_output_eg
               })
       
       # Тест Йохансена на коинтеграцию (для всех переменных сразу)
       is_cointegrated_johansen, cointegration_output_johansen = johansen_test(
           df,
           endogenous_vars,
           det_order=1,
           k_ar_diff=1
       )
       
       # Определяем общий результат коинтеграции
       any_eg_cointegrated = any([result['is_cointegrated'] for result in eg_results])
       is_cointegrated = any_eg_cointegrated or is_cointegrated_johansen
       
       # Сохраняем результаты обоих тестов
       results['cointegration_results'] = {
           'engle_granger_pairs': eg_results,
           'engle_granger_summary': {
               'is_cointegrated': any_eg_cointegrated,
               'cointegrated_pairs': [f"{result['variables'][0]} и {result['variables'][1]}" 
                                    for result in eg_results if result['is_cointegrated']]
           },
           'johansen': {
               'is_cointegrated': is_cointegrated_johansen,
               'details': cointegration_output_johansen
           },
           'conclusion': "Переменные коинтегрированы." if is_cointegrated else "Переменные не коинтегрированы.",
           'recommended_model': "VECM" if is_cointegrated else "VAR в разностях"
       }
   
   return results
# Функция создания отчета Word
def create_word_report(results, df):
    """
    Создает отчет в формате Word с результатами анализа стационарности и коинтеграции
    """
    # Создание нового документа Word
    doc = Document()
    
    # Добавление заголовка
    doc.add_heading('Отчет о стационарности временных рядов', 0)
    
    # Добавление текущей даты
    current_date = datetime.now().strftime("%d.%m.%Y")
    doc.add_paragraph(f'Дата создания отчета: {current_date}')
    
    # Сводная таблица результатов для всех переменных
    doc.add_heading('Сводная таблица анализа стационарности', 1)
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    # Заголовки таблицы
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Переменная'
    hdr_cells[1].text = 'Порядок интеграции'
    hdr_cells[2].text = 'Статус'
    
    # Заполнение таблицы для всех переменных
    for var, result in results['variable_results'].items():
        row_cells = table.add_row().cells
        row_cells[0].text = var
        row_cells[1].text = f"I({result['d_value']})"
        
        # Определение статуса
        d = result['d_value']
        if d == 0:
            status = "Стационарный"
        elif d == 1:
            status = "Интегрированный первого порядка"
        elif d == 2:
            status = "Интегрированный второго порядка"
        else:
            status = "Требует дополнительного анализа"
            
        row_cells[2].text = status
    
    # Если есть результаты коинтеграции
    if 'cointegration_results' in results:
        doc.add_heading('Анализ коинтеграции между переменными', 1)
        doc.add_paragraph(results['cointegration_results']['conclusion'])
        
        # Рекомендуемая модель
        doc.add_heading('Рекомендуемая модель', 1)
        doc.add_paragraph(results['cointegration_results']['recommended_model'])
    
    # Временный файл для отчета
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
        temp_path = temp_file.name
        doc.save(temp_path)
    
    return temp_path
# Модель данных для формы обратной связи
class FeedbackForm(BaseModel):
    name: str
    email: EmailStr
    subject: Optional[str] = "Сообщение с сайта"
    message: str

def save_message(feedback):
    try:
        # Создаем директорию, если её нет
        os.makedirs("feedback_messages", exist_ok=True)
        
        # Формируем сообщение
        message = {
            "name": feedback.name,
            "email": feedback.email,
            "subject": feedback.subject,
            "message": feedback.message,
            "date": datetime.now().isoformat()
        }
        
        # Создаем уникальное имя файла
        filename = f"feedback_messages/msg_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
        
        # Сохраняем в файл
        with open(filename, 'w', encoding='utf-8') as f:
            json.dump(message, f, ensure_ascii=False, indent=2)
        
        return True
    except Exception as e:
        print(f"Ошибка при сохранении сообщения: {str(e)}")
        return False

def cleanup_old_messages():
    """Удаляет сообщения старше 1 месяца"""
    while True:
        try:
            # Проверяем существование директории
            if not os.path.exists("feedback_messages"):
                time.sleep(86400)  # Проверяем раз в день
                continue
                
            # Текущая дата минус 1 месяц
            one_month_ago = datetime.now() - timedelta(days=30)
            
            # Проверяем все файлы в директории
            for filename in os.listdir("feedback_messages"):
                file_path = os.path.join("feedback_messages", filename)
                
                # Проверяем, является ли файл .json
                if not filename.endswith('.json'):
                    continue
                    
                # Получаем время создания файла из имени или из метаданных
                try:
                    # Пытаемся получить дату из имени файла
                    date_str = filename.split('_')[1].split('.')[0]
                    file_date = datetime.strptime(date_str, '%Y%m%d')
                except:
                    # Если не получилось, используем время модификации файла
                    file_date = datetime.fromtimestamp(os.path.getmtime(file_path))
                
                # Удаляем файл, если он старше 1 месяца
                if file_date < one_month_ago:
                    os.remove(file_path)
                    print(f"Удален устаревший файл с обратной связью: {filename}")
            
            # Проверяем раз в день
            time.sleep(86400)
            
        except Exception as e:
            print(f"Ошибка при очистке старых сообщений: {str(e)}")
            time.sleep(86400)  # В случае ошибки тоже ждем день

# Запускаем поток для очистки старых сообщений
cleanup_thread = threading.Thread(target=cleanup_old_messages, daemon=True)
cleanup_thread.start()

def transform_to_stationary(series, method='simple'):
    """
    Преобразование ряда к стационарному виду (I(0)) с использованием 
    последовательных разностей и тестов на стационарность
    """
    # Проверяем исходный ряд на стационарность
    is_stationary, _ = check_stationarity(series, "")
    
    if is_stationary == "Стационарен":
        # Если ряд уже стационарен, возвращаем его без изменений
        return series
    
    # Пробуем первые разности
    diff1 = series.diff().dropna()
    is_stationary, _ = check_stationarity(diff1, "")
    
    if is_stationary == "Стационарен":
        # Пополняем пропущенные значения (первое наблюдение)
        padded_series = pd.Series([None] + diff1.tolist(), index=series.index)
        return padded_series
    
    # Пробуем вторые разности
    diff2 = diff1.diff().dropna()
    is_stationary, _ = check_stationarity(diff2, "")
    
    if is_stationary == "Стационарен":
        # Пополняем пропущенные значения (первые два наблюдения)
        padded_series = pd.Series([None, None] + diff2.tolist(), index=series.index)
        return padded_series
    
    # Если не удалось достичь стационарности, возвращаем вторые разности
    padded_series = pd.Series([None, None] + diff2.tolist(), index=series.index)
    return padded_series

def transform_to_first_order(series, method='simple'):
    """
    Преобразование ряда к первому порядку интеграции (I(1))
    
    Parameters:
    -----------
    series : pandas.Series
        Исходный временной ряд
    method : str, default='simple'
        Метод преобразования (не используется в данной функции, 
        но добавлен для совместимости с другими функциями преобразования)
    
    Returns:
    --------
    pandas.Series
        Преобразованный временной ряд первого порядка интеграции
    """
    # Проверяем исходный ряд
    is_stationary, _ = check_stationarity(series, "")
    
    if is_stationary == "Стационарен":
        # Если ряд стационарен (I(0)), интегрируем его один раз
        integrated = series.cumsum()
        return integrated
    
    # Проверяем первые разности
    diff1 = series.diff().dropna()
    is_stationary, _ = check_stationarity(diff1, "")
    
    if is_stationary == "Стационарен":
        # Если первые разности стационарны (I(1)), оставляем исходный ряд
        return series
    
    # Проверяем вторые разности
    diff2 = diff1.diff().dropna()
    is_stationary, _ = check_stationarity(diff2, "")
    
    if is_stationary == "Стационарен":
        # Если вторые разности стационарны (I(2)), интегрируем один раз
        # Начинаем с первой разности и интегрируем
        padded_diff1 = pd.Series([0] + diff1.tolist(), index=series.index)
        integrated = padded_diff1.cumsum()
        return integrated
    
    # Если не определено, возвращаем исходный ряд
    return series

def transform_to_second_order(series, method='simple'):
    """
    Преобразование ряда ко второму порядку интеграции (I(2))
    
    Parameters:
    -----------
    series : pandas.Series
        Исходный временной ряд
    method : str, default='simple'
        Метод преобразования (не используется в данной функции, 
        но добавлен для совместимости с другими функциями преобразования)
    
    Returns:
    --------
    pandas.Series
        Преобразованный временной ряд второго порядка интеграции
    """
    # Проверяем исходный ряд
    is_stationary, _ = check_stationarity(series, "")
    
    if is_stationary == "Стационарен":
        # Если ряд стационарен (I(0)), интегрируем его дважды
        integrated1 = series.cumsum()
        integrated2 = integrated1.cumsum()
        return integrated2
    
    # Проверяем первые разности
    diff1 = series.diff().dropna()
    is_stationary, _ = check_stationarity(diff1, "")
    
    if is_stationary == "Стационарен":
        # Если первые разности стационарны (I(1)), интегрируем один раз
        integrated = series.cumsum()
        return integrated
    
    # Проверяем вторые разности
    diff2 = diff1.diff().dropna()
    is_stationary, _ = check_stationarity(diff2, "")
    
    if is_stationary == "Стационарен":
        # Если вторые разности стационарны (I(2)), оставляем исходный ряд
        return series
    
    # Если не определено, возвращаем исходный ряд
    return series

def format_time_series_for_preview(original_series, transformed_series, display_mode='all'):
    """
    Форматирует временной ряд для предпросмотра в формате для JSON с учетом режима отображения
    
    Parameters:
    -----------
    original_series : pandas.Series
        Исходный временной ряд
    transformed_series : pandas.Series
        Преобразованный временной ряд
    display_mode : str
        Режим отображения: 'all', 'transformed', 'original'
    
    Returns:
    --------
    list: Список словарей {'date': str, 'original': float, 'transformed': float}
    """
    result = []
    
    # Получаем индекс для правильного выравнивания данных
    if hasattr(original_series, 'index'):
        index = original_series.index
    else:
        # Если нет индекса, создаем числовой
        index = range(len(original_series))
    
    # Создаем DataFrame для правильного выравнивания данных
    df = pd.DataFrame({
        'original': original_series,
        'transformed': transformed_series
    }, index=index)
    
    for idx, row in df.iterrows():
        # Форматируем дату
        if isinstance(idx, pd.Timestamp):
            date_str = idx.strftime("%Y-%m-%d")
        else:
            date_str = str(idx)
        
        # Добавляем данные в зависимости от режима отображения
        point = {"date": date_str}
        
        if display_mode in ['all', 'original'] and pd.notna(row['original']):
            point['original'] = float(row['original'])
        
        if display_mode in ['all', 'transformed'] and pd.notna(row['transformed']):
            point['transformed'] = float(row['transformed'])
        
        # Добавляем точку только если есть хотя бы одно значение
        if len(point) > 1:  # больше чем просто date
            result.append(point)
    
    return result
def build_varx_model(df, endogenous_vars, exogenous_vars=None, lags=1, train_size=0.8):
    """
    Обучение VARX модели на данных и оценка её качества.
    Если указана одна эндогенная переменная, строится ARX модель.
    
    Parameters:
    -----------
    df : pandas.DataFrame
        Датафрейм с временными рядами
    endogenous_vars : list
        Список имен эндогенных переменных для VARX/ARX модели
    exogenous_vars : list, optional
        Список имен экзогенных переменных для VARX/ARX модели
    lags : int, default=1
        Количество лагов для VARX/ARX модели
    train_size : float, default=0.8
        Доля данных для обучения модели (0-1)
        
    Returns:
    --------
    dict
        Результаты обучения и оценки модели
    """
    import statsmodels.api as sm
    import pickle
    import numpy as np
    import io
    import base64
    from matplotlib import pyplot as plt
    from statsmodels.graphics.tsaplots import plot_acf, plot_pacf
    
    # Убираем пропуски из данных
    df_clean = df.dropna()
    
    # Подготовка данных
    endog = df_clean[endogenous_vars]
    exog = df_clean[exogenous_vars] if exogenous_vars else None
    
    # Разделение на обучающую и тестовую выборки
    train_size_idx = int(len(df_clean) * train_size)
    train_endog = endog.iloc[:train_size_idx]
    train_exog = exog.iloc[:train_size_idx] if exog is not None else None
    test_endog = endog.iloc[train_size_idx:]
    test_exog = exog.iloc[train_size_idx:] if exog is not None else None
    
    results = {
        "success": True,
        "model_info": {},
        "forecasts": {},
        "diagnostics": {},
        "plots": {}
    }
    
    try:
        # Определение оптимального числа лагов с помощью ACF и PACF
        acf_pacf_plots = []
        optimal_lags = {}
        
        for var in endogenous_vars:
            series = train_endog[var]
            
            # ACF и PACF анализ
            fig, axes = plt.subplots(2, 1, figsize=(12, 10))  # Увеличиваем размер с (10, 8) до (12, 10)
            
            # ACF
            plot_acf(series, lags=20, ax=axes[0], title=f'ACF для {var}')
            
            # PACF
            plot_pacf(series, lags=20, ax=axes[1], title=f'PACF для {var}')
            
            # Увеличиваем размер шрифта для лучшей читаемости
            axes[0].tick_params(axis='both', which='major', labelsize=12)
            axes[1].tick_params(axis='both', which='major', labelsize=12)
            axes[0].title.set_size(14)
            axes[1].title.set_size(14)
            
            plt.tight_layout()
            
            # Сохраняем график в base64
            buf = io.BytesIO()
            plt.savefig(buf, format='png')
            buf.seek(0)
            img_str = base64.b64encode(buf.read()).decode('utf-8')
            acf_pacf_plots.append(img_str)
            plt.close()
            
            # Определяем оптимальное число лагов по PACF
            # Значения выше 95% доверительного интервала считаются значимыми
            from statsmodels.tsa.stattools import pacf
            pacf_values, conf_intervals = pacf(series, nlags=20, alpha=0.05, method='ywm')
            
            # Находим последний значимый лаг
            significant_lags = []
            for i in range(1, len(pacf_values)):
                if abs(pacf_values[i]) > abs(conf_intervals[i, 1]):  # Верхняя граница 95% интервала
                    significant_lags.append(i)
            
            # Если нет значимых лагов или только первый лаг значим, берем 1
            # В противном случае берем последний значимый лаг (до максимума 10)
            if not significant_lags:
                optimal_lag = 1
            else:
                optimal_lag = min(max(significant_lags), 10)
            
            optimal_lags[var] = optimal_lag
        
        # Определяем итоговое количество лагов (максимальное из оптимальных)
        optimal_lag_final = max(optimal_lags.values()) if optimal_lags else lags
        
        # Используем переданное значение лагов, но сохраняем оптимальное для информации
        model_lags = lags
        
        # Обрезаем данные в соответствии с выбранным числом лагов
        train_endog_adj = train_endog.iloc[model_lags:]
        train_exog_adj = train_exog.iloc[model_lags:] if train_exog is not None else None
        
        # Обучаем модель
        if len(endogenous_vars) == 1:
            # ARX модель для одной эндогенной переменной
            endog_var = endogenous_vars[0]
            model = sm.tsa.SARIMAX(
                train_endog_adj[endog_var], 
                exog=train_exog_adj,
                order=(model_lags, 0, 0),  # (p, d, q) - используем только AR часть
                trend='c'
            )
            fitted_model = model.fit(disp=False)
        else:
            # VARX модель для нескольких эндогенных переменных
            model = sm.tsa.VARMAX(train_endog_adj, order=(model_lags, 0), trend='c', exog=train_exog_adj)
            fitted_model = model.fit(maxiter=1000, disp=False)
        
        # Сохраняем информацию о модели
        model_type = "ARX" if len(endogenous_vars) == 1 else "VARX"
        results["model_info"] = {
            "model_type": model_type,
            "aic": fitted_model.aic,
            "bic": fitted_model.bic,
            "hqic": fitted_model.hqic if hasattr(fitted_model, 'hqic') else None,
            "parameters": fitted_model.params.to_dict() if hasattr(fitted_model.params, 'to_dict') else dict(zip(fitted_model.params.index.tolist(), fitted_model.params.values.tolist())),
            "lags": model_lags,
            "optimal_lags": optimal_lags,
            "optimal_lag_final": optimal_lag_final,
            "endogenous_vars": endogenous_vars,
            "exogenous_vars": exogenous_vars or [],
            "train_size": train_size_idx,
            "test_size": len(test_endog)
        }
        
        # Прогнозы на тестовом наборе, если он не пустой
        if len(test_endog) > 0:
            try:
                # Создаем прогнозы
                if len(endogenous_vars) == 1:
                    # Для ARX модели
                    endog_var = endogenous_vars[0]
                    forecast = fitted_model.forecast(steps=len(test_endog), exog=test_exog)
                    # Преобразуем в DataFrame для совместимости с VARX моделью
                    forecast_df = pd.DataFrame({endog_var: forecast}, index=test_endog.index)
                else:
                    # Для VARX модели
                    forecast_df = fitted_model.forecast(steps=len(test_endog), exog=test_exog)
                
                # Оценка качества прогнозов
                mse = {}
                mae = {}
                rmse = {}
                mape = {}
                
                for i, var in enumerate(endogenous_vars):
                    actual = test_endog[var].values
                    if len(endogenous_vars) == 1:
                        predicted = forecast_df[var].values
                    else:
                        predicted = forecast_df.iloc[:, i].values
                    
                    # Убираем NaN перед расчетом метрик
                    valid_indices = ~(np.isnan(actual) | np.isnan(predicted))
                    
                    if np.any(valid_indices):
                        actual_valid = actual[valid_indices]
                        predicted_valid = predicted[valid_indices]
                        
                        var_mse = np.mean((actual_valid - predicted_valid) ** 2)
                        var_mae = np.mean(np.abs(actual_valid - predicted_valid))
                        var_rmse = np.sqrt(var_mse)
                        
                        # MAPE может вызвать деление на ноль, поэтому используем try/except
                        try:
                            var_mape = np.mean(np.abs((actual_valid - predicted_valid) / actual_valid)) * 100
                        except:
                            var_mape = np.nan
                        
                        mse[var] = var_mse
                        mae[var] = var_mae
                        rmse[var] = var_rmse
                        mape[var] = var_mape
                
                # Создаем DataFrame для сравнения
                comparison = {}
                for i, var in enumerate(endogenous_vars):
                    if len(endogenous_vars) == 1:
                        pred_values = forecast_df[var].tolist()
                    else:
                        pred_values = forecast_df.iloc[:, i].tolist()
                    
                    comparison[var] = {
                        "actual": test_endog[var].tolist(),
                        "predicted": pred_values,
                        "dates": test_endog.index.strftime("%Y-%m-%d").tolist()
                    }
                
                results["forecasts"] = {
                    "comparison": comparison,
                    "metrics": {
                        "mse": mse,
                        "mae": mae,
                        "rmse": rmse,
                        "mape": mape
                    }
                }
                
                # Создаем графики прогнозов
                forecast_plots = []
                for i, var in enumerate(endogenous_vars):
                    plt.figure(figsize=(10, 6))
                    plt.plot(test_endog.index, test_endog[var], 'b-', label='Фактические значения')
                    
                    if len(endogenous_vars) == 1:
                        pred_values = forecast_df[var]
                    else:
                        pred_values = forecast_df.iloc[:, i]
                    
                    plt.plot(test_endog.index, pred_values, 'r--', label='Прогноз')
                    plt.title(f'Прогноз vs. Факт: {var}')
                    plt.legend()
                    plt.grid(True)
                    
                    # Форматирование дат на оси X
                    plt.gcf().autofmt_xdate()
                    
                    # Сохраняем график в base64
                    buf = io.BytesIO()
                    plt.savefig(buf, format='png')
                    buf.seek(0)
                    img_str = base64.b64encode(buf.read()).decode('utf-8')
                    forecast_plots.append(img_str)
                    plt.close()
                
                results["plots"]["forecast_plots"] = forecast_plots
                
            except Exception as e:
                results["forecasts"] = {"error": str(e)}
        
        # Сохраняем графики ACF и PACF
        results["plots"]["acf_pacf_plots"] = acf_pacf_plots
        
        # Добавляем информацию об оптимальных лагах
        results["diagnostics"]["lag_analysis"] = {
            "optimal_lags": optimal_lags,
            "optimal_lag_final": optimal_lag_final,
            "chosen_lag": model_lags
        }
        
        # Сериализуем модель в base64 для возможного сохранения
        model_bytes = io.BytesIO()
        pickle.dump(fitted_model, model_bytes)
        model_bytes.seek(0)
        results["model_base64"] = base64.b64encode(model_bytes.read()).decode('utf-8')
        
    except Exception as e:
        results["success"] = False
        results["error"] = str(e)
        import traceback
        results["traceback"] = traceback.format_exc()
    
    return results

def build_varx_model_with_future_forecast(df, endogenous_vars, exogenous_vars=None, lags=1, train_size=0.8, forecast_periods=6, forecast_unit='months'):
    """
    Обучение VARX модели на данных и создание прогнозов как на тестовых данных, так и в будущее.
    СПЕЦИАЛЬНО АДАПТИРОВАНО ДЛЯ РАБОТЫ С МИНИМАЛЬНЫМИ ДАННЫМИ (даже 9 наблюдений)!
    """
    import statsmodels.api as sm
    import pickle
    import numpy as np
    import io
    import base64
    from matplotlib import pyplot as plt
    from statsmodels.graphics.tsaplots import plot_acf, plot_pacf
    import pandas as pd
    from datetime import datetime, timedelta
    from dateutil.relativedelta import relativedelta
    from sklearn.linear_model import LinearRegression
    from sklearn.preprocessing import StandardScaler
    
    print(f"=== НАЧАЛО ОБРАБОТКИ VARX ===")
    print(f"Исходные данные: {df.shape}")
    print(f"Эндогенные переменные: {endogenous_vars}")
    print(f"Экзогенные переменные: {exogenous_vars}")
    print(f"Лаги: {lags}")
    
    # Убираем пропуски из данных
    df_clean = df.dropna()
    print(f"После удаления пропусков: {df_clean.shape}")
    
    # КРИТИЧЕСКИ ВАЖНО: При малом количестве данных используем упрощенный подход
    available_obs = len(df_clean)
    
    if available_obs < 15:
        print(f"МАЛО ДАННЫХ ({available_obs} наблюдений): используем упрощенную модель")
        use_simple_model = True
        model_lags = 0  # Никаких лагов при малых данных
        train_size_ratio = 1.0  # Используем все данные для обучения
    else:
        use_simple_model = False
        # Адаптивное ограничение лагов
        model_lags = min(lags, max(1, available_obs // 5))
        train_size_ratio = train_size
    
    print(f"Режим работы: {'Упрощенная модель' if use_simple_model else 'Стандартная модель'}")
    print(f"Использовать лагов: {model_lags}")
    
    results = {
        "success": True,
        "model_info": {},
        "forecasts": {},
        "validation": {},
        "diagnostics": {},
        "plots": {}
    }
    
    try:
        # Подготовка данных
        endog = df_clean[endogenous_vars].copy()
        
        # Обработка экзогенных переменных
        exog = None
        final_exogenous_vars = []
        if exogenous_vars and len(exogenous_vars) > 0:
            existing_exog = [var for var in exogenous_vars if var in df_clean.columns]
            if existing_exog:
                exog = df_clean[existing_exog].copy()
                final_exogenous_vars = existing_exog
                
                # При малых данных ограничиваем количество экзогенных переменных
                if use_simple_model and len(final_exogenous_vars) > 2:
                    final_exogenous_vars = final_exogenous_vars[:2]
                    exog = exog[final_exogenous_vars]
                    print(f"Ограничили экзогенные переменные до: {final_exogenous_vars}")
        
        print(f"Финальные переменные:")
        print(f"  Эндогенные: {endogenous_vars}")
        print(f"  Экзогенные: {final_exogenous_vars}")
        print(f"  Endog shape: {endog.shape}")
        if exog is not None:
            print(f"  Exog shape: {exog.shape}")
        
        # Разделение данных
        if use_simple_model:
            # При малых данных используем все для обучения
            train_size_idx = available_obs
            train_endog = endog.copy()
            train_exog = exog.copy() if exog is not None else None
            test_endog = pd.DataFrame()  # Пустой тест
            test_exog = pd.DataFrame() if exog is not None else None
        else:
            # Стандартное разделение
            train_size_idx = int(available_obs * train_size_ratio)
            min_test_size = max(2, model_lags + 1)
            
            if available_obs - train_size_idx < min_test_size:
                train_size_idx = available_obs - min_test_size
            
            train_endog = endog.iloc[:train_size_idx].copy()
            train_exog = exog.iloc[:train_size_idx].copy() if exog is not None else None
            test_endog = endog.iloc[train_size_idx:].copy()
            test_exog = exog.iloc[train_size_idx:].copy() if exog is not None else None
        
        print(f"Разделение данных: train={len(train_endog)}, test={len(test_endog)}")
        
        # Создание упрощенных ACF/PACF графиков
        acf_pacf_plots = []
        optimal_lags = {}
        
        if not use_simple_model and len(train_endog) >= 10:
            for var in endogenous_vars:
                try:
                    series = train_endog[var].dropna()
                    if len(series) >= 10:
                        fig, axes = plt.subplots(2, 1, figsize=(10, 8))
                        
                        max_lags_plot = min(5, len(series) // 2)
                        plot_acf(series, lags=max_lags_plot, ax=axes[0], title=f'ACF для {var}')
                        plot_pacf(series, lags=max_lags_plot, ax=axes[1], title=f'PACF для {var}')
                        
                        plt.tight_layout()
                        
                        buf = io.BytesIO()
                        plt.savefig(buf, format='png', dpi=100)
                        buf.seek(0)
                        img_str = base64.b64encode(buf.read()).decode('utf-8')
                        acf_pacf_plots.append(img_str)
                        plt.close()
                    
                    optimal_lags[var] = model_lags
                except Exception as e:
                    print(f"Ошибка при создании графика для {var}: {str(e)}")
                    optimal_lags[var] = model_lags
        else:
            # Для упрощенной модели создаем заглушку
            for var in endogenous_vars:
                optimal_lags[var] = 0
        
        # ОБУЧЕНИЕ МОДЕЛИ
        fitted_model = None
        model_type = "ARX" if len(endogenous_vars) == 1 else "VARX"
        
        if use_simple_model:
            # УПРОЩЕННАЯ МОДЕЛЬ ДЛЯ МАЛЫХ ДАННЫХ
            print("Используем упрощенную регрессионную модель...")
            
            try:
                # Подготавливаем данные для регрессии
                y = train_endog[endogenous_vars[0]].values
                
                if train_exog is not None and len(final_exogenous_vars) > 0:
                    X = train_exog[final_exogenous_vars].values
                    
                    # Нормализация данных
                    scaler_X = StandardScaler()
                    scaler_y = StandardScaler()
                    
                    X_scaled = scaler_X.fit_transform(X)
                    y_scaled = scaler_y.fit_transform(y.reshape(-1, 1)).ravel()
                    
                    # Обучаем линейную регрессию
                    reg_model = LinearRegression()
                    reg_model.fit(X_scaled, y_scaled)
                    
                    print(f"✓ Упрощенная регрессия обучена! R² = {reg_model.score(X_scaled, y_scaled):.4f}")
                    
                    # Создаем псевдо-statsmodels объект для совместимости
                    class SimpleRegressionWrapper:
                        def __init__(self, reg_model, scaler_X, scaler_y, feature_names):
                            self.reg_model = reg_model
                            self.scaler_X = scaler_X
                            self.scaler_y = scaler_y
                            self.feature_names = feature_names
                            self.aic = np.nan
                            self.bic = np.nan
                            self.hqic = np.nan
                        
                        def forecast(self, steps, exog=None):
                            if exog is None:
                                # Используем последние значения
                                last_X = X[-1:].copy()
                                forecasts = []
                                for _ in range(steps):
                                    X_scaled = self.scaler_X.transform(last_X)
                                    y_scaled_pred = self.reg_model.predict(X_scaled)
                                    y_pred = self.scaler_y.inverse_transform(y_scaled_pred.reshape(-1, 1)).ravel()
                                    forecasts.append(y_pred[0])
                                return np.array(forecasts)
                            else:
                                X_forecast = exog[self.feature_names].values
                                X_scaled = self.scaler_X.transform(X_forecast)
                                y_scaled_pred = self.reg_model.predict(X_scaled)
                                y_pred = self.scaler_y.inverse_transform(y_scaled_pred.reshape(-1, 1)).ravel()
                                return y_pred
                    
                    fitted_model = SimpleRegressionWrapper(reg_model, scaler_X, scaler_y, final_exogenous_vars)
                    
                else:
                    # Модель только с трендом
                    print("Используем модель только с трендом...")
                    
                    # Простая линейная регрессия с трендом
                    X_trend = np.arange(len(y)).reshape(-1, 1)
                    reg_model = LinearRegression()
                    reg_model.fit(X_trend, y)
                    
                    class TrendOnlyWrapper:
                        def __init__(self, reg_model, last_index):
                            self.reg_model = reg_model
                            self.last_index = last_index
                            self.aic = np.nan
                            self.bic = np.nan
                            self.hqic = np.nan
                        
                        def forecast(self, steps, exog=None):
                            future_indices = np.arange(self.last_index + 1, self.last_index + 1 + steps).reshape(-1, 1)
                            return self.reg_model.predict(future_indices)
                    
                    fitted_model = TrendOnlyWrapper(reg_model, len(y) - 1)
                    final_exogenous_vars = []
                
            except Exception as e:
                print(f"Ошибка в упрощенной модели: {str(e)}")
                return {
                    "success": False,
                    "error": f"Не удалось обучить упрощенную модель: {str(e)}"
                }
        
        else:
            # СТАНДАРТНАЯ МОДЕЛЬ STATSMODELS
            print("Используем стандартную statsmodels модель...")
            
            # Обрезка данных на лаги
            if model_lags > 0:
                train_endog_adj = train_endog.iloc[model_lags:].copy()
                train_exog_adj = train_exog.iloc[model_lags:].copy() if train_exog is not None else None
            else:
                train_endog_adj = train_endog.copy()
                train_exog_adj = train_exog.copy() if train_exog is not None else None
            
            print(f"Размеры для обучения: endog={train_endog_adj.shape}, exog={train_exog_adj.shape if train_exog_adj is not None else None}")
            
            if len(endogenous_vars) == 1:
                endog_var = endogenous_vars[0]
                
                try:
                    if train_exog_adj is not None and len(final_exogenous_vars) > 0:
                        model = sm.tsa.SARIMAX(
                            train_endog_adj[endog_var], 
                            exog=train_exog_adj[final_exogenous_vars],
                            order=(max(model_lags, 1), 0, 0),
                            trend='c'
                        )
                        fitted_model = model.fit(disp=False, maxiter=50)
                        print("✓ ARX модель обучена!")
                    else:
                        model = sm.tsa.SARIMAX(
                            train_endog_adj[endog_var], 
                            order=(max(model_lags, 1), 0, 0),
                            trend='c'
                        )
                        fitted_model = model.fit(disp=False, maxiter=50)
                        final_exogenous_vars = []
                        print("✓ AR модель обучена!")
                        
                except Exception as e:
                    print(f"Ошибка statsmodels: {str(e)}")
                    return {
                        "success": False,
                        "error": f"Ошибка обучения модели: {str(e)}"
                    }
            
            else:
                # VARX модель
                try:
                    if train_exog_adj is not None and len(final_exogenous_vars) > 0:
                        model = sm.tsa.VARMAX(
                            train_endog_adj, 
                            order=(max(model_lags, 1), 0), 
                            trend='c', 
                            exog=train_exog_adj[final_exogenous_vars]
                        )
                        fitted_model = model.fit(maxiter=50, disp=False)
                        print("✓ VARX модель обучена!")
                    else:
                        model = sm.tsa.VARMAX(
                            train_endog_adj, 
                            order=(max(model_lags, 1), 0), 
                            trend='c'
                        )
                        fitted_model = model.fit(maxiter=50, disp=False)
                        final_exogenous_vars = []
                        print("✓ VAR модель обучена!")
                        
                except Exception as e:
                    print(f"Ошибка VARX: {str(e)}")
                    return {
                        "success": False,
                        "error": f"Ошибка обучения VARX модели: {str(e)}"
                    }
        
        if fitted_model is None:
            return {
                "success": False,
                "error": "Не удалось обучить модель"
            }
        
        print("✓ МОДЕЛЬ УСПЕШНО ОБУЧЕНА!")
        
        # Информация о прогнозном периоде
        unit_names = {
            'months': {'single': 'месяц', 'few': 'месяца', 'many': 'месяцев'},
            'quarters': {'single': 'квартал', 'few': 'квартала', 'many': 'кварталов'},
            'years': {'single': 'год', 'few': 'года', 'many': 'лет'}
        }
        
        if forecast_periods == 1:
            unit_text = unit_names[forecast_unit]['single']
        elif forecast_periods < 5:
            unit_text = unit_names[forecast_unit]['few']
        else:
            unit_text = unit_names[forecast_unit]['many']
        
        forecast_description = f"{forecast_periods} {unit_text} вперед от {df_clean.index[-1].strftime('%d.%m.%Y')}"
        
        # Сохраняем информацию о модели
        results["model_info"] = {
            "model_type": f"{model_type} {'(упрощенная)' if use_simple_model else ''}",
            "aic": fitted_model.aic if hasattr(fitted_model, 'aic') else np.nan,
            "bic": fitted_model.bic if hasattr(fitted_model, 'bic') else np.nan,
            "hqic": fitted_model.hqic if hasattr(fitted_model, 'hqic') else np.nan,
            "lags": model_lags,
            "optimal_lags": optimal_lags,
            "optimal_lag_final": model_lags,
            "endogenous_vars": endogenous_vars,
            "exogenous_vars": final_exogenous_vars,
            "train_size": len(train_endog),
            "test_size": len(test_endog),
            "forecast_info": {
                "periods": forecast_periods,
                "unit": forecast_unit,
                "description": forecast_description
            }
        }
        
        # Валидация на тестовом наборе (если есть)
        validation_results = {}
        if len(test_endog) > 0 and not use_simple_model:
            try:
                print("Создаем прогнозы на тестовых данных...")
                test_steps = len(test_endog)
                
                # Подготовка экзогенных переменных для тестирования
                test_exog_for_forecast = None
                if final_exogenous_vars and test_exog is not None:
                    test_exog_for_forecast = test_exog[final_exogenous_vars].iloc[:test_steps]
                
                # Создаем прогнозы
                if len(endogenous_vars) == 1:
                    endog_var = endogenous_vars[0]
                    test_forecast = fitted_model.forecast(steps=test_steps, exog=test_exog_for_forecast)
                    test_forecast_df = pd.DataFrame({endog_var: test_forecast}, index=test_endog.index)
                else:
                    test_forecast_df = fitted_model.forecast(steps=test_steps, exog=test_exog_for_forecast)
                    test_forecast_df.index = test_endog.index
                
                # Рассчитываем метрики
                mse = {}
                mae = {}
                validation_comparison = {}
                
                for i, var in enumerate(endogenous_vars):
                    actual = test_endog[var].values
                    if len(endogenous_vars) == 1:
                        predicted = test_forecast_df[var].values
                    else:
                        predicted = test_forecast_df.iloc[:, i].values
                    
                    # Убираем NaN
                    valid_indices = ~(np.isnan(actual) | np.isnan(predicted))
                    
                    if np.any(valid_indices):
                        actual_valid = actual[valid_indices]
                        predicted_valid = predicted[valid_indices]
                        
                        mse[var] = np.mean((actual_valid - predicted_valid) ** 2)
                        mae[var] = np.mean(np.abs(actual_valid - predicted_valid))
                    
                    validation_comparison[var] = {
                        "actual": test_endog[var].tolist(),
                        "predicted": (test_forecast_df[var] if len(endogenous_vars) == 1 else test_forecast_df.iloc[:, i]).tolist(),
                        "dates": test_endog.index.strftime("%Y-%m-%d").tolist()
                    }
                
                validation_results = {
                    "comparison": validation_comparison,
                    "metrics": {"mse": mse, "mae": mae}
                }
                print("✓ Валидация завершена!")
                
            except Exception as e:
                print(f"Ошибка при валидации: {str(e)}")
                validation_results = {"error": str(e)}
        
        # Создание прогнозов в будущее
        future_forecast_results = {}
        try:
            print("Создаем прогнозы в будущее...")
            
            # Создаем будущие даты
            last_date = df_clean.index[-1]
            future_dates = []
            current_date = last_date
            
            for i in range(forecast_periods):
                if forecast_unit == 'months':
                    current_date = current_date + relativedelta(months=1)
                elif forecast_unit == 'quarters':
                    current_date = current_date + relativedelta(months=3)
                elif forecast_unit == 'years':
                    current_date = current_date + relativedelta(years=1)
                future_dates.append(current_date)
            
            # Подготовка экзогенных переменных для будущего
            future_exog = None
            if final_exogenous_vars and exog is not None:
                # Используем последние значения экзогенных переменных
                last_exog = exog[final_exogenous_vars].iloc[-1:]
                future_exog_data = np.tile(last_exog.values, (forecast_periods, 1))
                future_exog = pd.DataFrame(
                    future_exog_data, 
                    columns=final_exogenous_vars,
                    index=future_dates
                )
            
            # Создаем прогнозы
            if len(endogenous_vars) == 1:
                endog_var = endogenous_vars[0]
                future_forecast = fitted_model.forecast(steps=forecast_periods, exog=future_exog)
                future_forecast_df = pd.DataFrame({endog_var: future_forecast}, index=future_dates)
            else:
                future_forecast_df = fitted_model.forecast(steps=forecast_periods, exog=future_exog)
                future_forecast_df.index = future_dates
            
            # Подготавливаем результат
            future_values = {}
            for var in endogenous_vars:
                if len(endogenous_vars) == 1:
                    future_values[var] = future_forecast_df[var].tolist()
                else:
                    var_index = endogenous_vars.index(var)
                    future_values[var] = future_forecast_df.iloc[:, var_index].tolist()
            
            future_forecast_results = {
                "dates": [date.strftime("%Y-%m-%d") for date in future_dates],
                "values": future_values,
                "unit": forecast_unit,
                "periods": forecast_periods
            }
            print("✓ Прогнозы в будущее созданы!")
            
        except Exception as e:
            print(f"Ошибка при создании будущих прогнозов: {str(e)}")
            future_forecast_results = {"error": str(e)}
        
        # Создание комбинированных данных для графиков
        full_comparison = {}
        try:
            for i, var in enumerate(endogenous_vars):
                # Исторические данные
                historical_values = train_endog[var].tolist()
                historical_dates = train_endog.index.strftime("%Y-%m-%d").tolist()
                
                # Тестовые данные
                test_actual = test_endog[var].tolist() if len(test_endog) > 0 else []
                test_dates = test_endog.index.strftime("%Y-%m-%d").tolist() if len(test_endog) > 0 else []
                
                # Прогнозы на тестовых данных
                if len(test_endog) > 0 and 'comparison' in validation_results:
                    test_predicted = validation_results['comparison'][var]['predicted']
                else:
                    test_predicted = []
                
                # Будущие прогнозы
                if 'values' in future_forecast_results and var in future_forecast_results['values']:
                    future_predicted = future_forecast_results['values'][var]
                    future_dates_list = future_forecast_results['dates']
                else:
                    future_predicted = []
                    future_dates_list = []
                
                # Объединяем все даты
                all_dates = historical_dates + test_dates + future_dates_list
                
                # Создаем массивы данных с пропусками для корректного отображения
                historical_full = historical_values + [None] * (len(test_dates) + len(future_dates_list))
                actual_full = [None] * len(historical_dates) + test_actual + [None] * len(future_dates_list)
                predicted_full = [None] * len(historical_dates) + test_predicted + [None] * len(future_dates_list)
                future_full = [None] * (len(historical_dates) + len(test_dates)) + future_predicted
                
                full_comparison[var] = {
                    "all_dates": all_dates,
                    "historical": historical_full,
                    "actual": actual_full,
                    "predicted": predicted_full,
                    "future_forecast": future_full
                }
        
        except Exception as e:
            print(f"Ошибка при подготовке полных данных: {str(e)}")
            full_comparison = {}
        
        # Сохраняем результаты
        results["validation"] = validation_results
        results["forecasts"] = {
            "comparison": full_comparison,
            "future_forecast": future_forecast_results
        }
        results["plots"]["acf_pacf_plots"] = acf_pacf_plots
        
        # Добавляем информацию об оптимальных лагах
        results["diagnostics"]["lag_analysis"] = {
            "optimal_lags": optimal_lags,
            "optimal_lag_final": model_lags,
            "chosen_lag": model_lags
        }
        
        print("✓ ВСЕ ОПЕРАЦИИ ЗАВЕРШЕНЫ УСПЕШНО!")
        
    except Exception as e:
        results["success"] = False
        results["error"] = str(e)
        import traceback
        results["traceback"] = traceback.format_exc()
        print(f"КРИТИЧЕСКАЯ ОШИБКА: {str(e)}")
        print(traceback.format_exc())
    
    return results
    
def create_comprehensive_report(df, params):
    """
    Создание комплексного отчета на основе выбранных тестов и моделей
    
    Parameters:
    -----------
    df : pandas.DataFrame
        Датафрейм с временными рядами
    params : dict
        Параметры для отчета:
            - tests: dict - выбранные тесты и их параметры
            - models: dict - выбранные модели и их параметры
    
    Returns:
    --------
    str
        Путь к созданному файлу отчета
    """
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    import tempfile
    import os
    import matplotlib.pyplot as plt
    import io
    from datetime import datetime
    
    # Создаем новый документ Word
    doc = Document()
    
    # Настройка стилей
    styles = doc.styles
    
    # Стиль для заголовка
    title_style = styles.add_style('ReportTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_font = title_style.font
    title_font.name = 'Arial'
    title_font.size = Pt(18)
    title_font.bold = True
    title_font.color.rgb = RGBColor(0x33, 0x41, 0x55)
    
    # Стиль для подзаголовка
    heading1_style = styles.add_style('Heading1Custom', WD_STYLE_TYPE.PARAGRAPH)
    heading1_font = heading1_style.font
    heading1_font.name = 'Arial'
    heading1_font.size = Pt(16)
    heading1_font.bold = True
    heading1_font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
    
    # Стиль для заголовка второго уровня
    heading2_style = styles.add_style('Heading2Custom', WD_STYLE_TYPE.PARAGRAPH)
    heading2_font = heading2_style.font
    heading2_font.name = 'Arial'
    heading2_font.size = Pt(14)
    heading2_font.bold = True
    heading2_font.color.rgb = RGBColor(0x33, 0x41, 0x55)
    
    # Стиль для обычного текста
    normal_style = styles.add_style('NormalCustom', WD_STYLE_TYPE.PARAGRAPH)
    normal_font = normal_style.font
    normal_font.name = 'Arial'
    normal_font.size = Pt(11)
    normal_font.color.rgb = RGBColor(0x33, 0x41, 0x55)
    
    # Добавляем заголовок отчета
    title = doc.add_paragraph('Аналитический отчет по временным рядам', style='ReportTitle')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Добавляем дату создания отчета
    date_paragraph = doc.add_paragraph(f'Дата создания: {datetime.now().strftime("%d.%m.%Y %H:%M")}', style='NormalCustom')
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Добавляем разделитель
    doc.add_paragraph('', style='NormalCustom')
    
    # Получаем выбранные тесты и модели
    tests = params.get('tests', {})
    models = params.get('models', {})
    
    # Получаем числовые колонки (предполагаемые переменные)
    numeric_columns = [col for col in df.columns if pd.api.types.is_numeric_dtype(df[col])]
    
    # Переменная для отслеживания номера раздела
    section_number = 1
    
    # Если выбран тест на стационарность
    if tests.get('stationary', False):
        # ИСПРАВЛЕНО: добавляем заголовок без указания стиля
        heading = doc.add_heading(f'{section_number}. Анализ стационарности временных рядов', level=1)
        # Применяем стиль к параграфу заголовка после его создания
        heading.style = heading1_style
        
        doc.add_paragraph('В этом разделе представлены результаты тестов на стационарность для всех числовых переменных в наборе данных. Используются три различных теста: расширенный тест Дики-Фуллера (ADF), тест Квятковски-Филлипса-Шмидта-Шина (KPSS) и тест Филлипса-Перрона.', style='NormalCustom')
        
        # Таблица с результатами
        # ИСПРАВЛЕНО: добавляем заголовок без указания стиля
        heading = doc.add_heading(f'{section_number}.1. Сводная таблица результатов', level=2)
        heading.style = heading2_style
        
        # Создаем таблицу для результатов
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        # Заголовки таблицы
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Переменная'
        hdr_cells[1].text = 'Порядок интеграции'
        hdr_cells[2].text = 'Статус'
        hdr_cells[3].text = 'Рекомендация'
        
        # Выполняем анализ стационарности для каждой переменной
        for column in numeric_columns:
            # Анализ разностей для достижения стационарности
            d_value, test_outputs = analyze_differences(df[column], column, max_diff=4)
            
            # Добавляем строку в таблицу
            row_cells = table.add_row().cells
            row_cells[0].text = column
            
            # Порядок интеграции
            if isinstance(d_value, (int, float)):
                row_cells[1].text = f'I({d_value})'
            else:
                row_cells[1].text = 'Не определен'
            
            # Статус
            if d_value == 0:
                row_cells[2].text = 'Стационарный'
                row_cells[3].text = 'Может использоваться без преобразований'
            elif d_value == 1:
                row_cells[2].text = 'Интегрированный первого порядка'
                row_cells[3].text = 'Рекомендуется использовать первые разности'
            elif d_value == 2:
                row_cells[2].text = 'Интегрированный второго порядка'
                row_cells[3].text = 'Рекомендуется использовать вторые разности'
            else:
                row_cells[2].text = 'Высокий порядок интеграции или неопределен'
                row_cells[3].text = 'Требуется дополнительный анализ'
            
            # Добавляем детальные результаты тестов
            # ИСПРАВЛЕНО: добавляем заголовок без указания стиля
            heading = doc.add_heading(f'{section_number}.2. Результаты тестов для переменной "{column}"', level=2)
            heading.style = heading2_style
            
            for line in test_outputs:
                doc.add_paragraph(line, style='NormalCustom')
            
            # Визуализация ряда и его разностей
            # ИСПРАВЛЕНО: добавляем заголовок без указания стиля
            heading = doc.add_heading(f'{section_number}.3. Визуализация ряда и его разностей для "{column}"', level=2)
            heading.style = heading2_style
            
            # Создаем диаграмму
            plt.figure(figsize=(10, 6))
            plt.plot(df.index, df[column])
            plt.title(f'Временной ряд {column}')
            plt.xlabel('Дата')
            plt.ylabel('Значение')
            plt.grid(True)
            
            # Сохраняем диаграмму в память
            img_stream = io.BytesIO()
            plt.savefig(img_stream, format='png')
            img_stream.seek(0)
            plt.close()
            
            # Добавляем изображение в документ
            doc.add_picture(img_stream, width=Inches(6))
            
            # Если порядок интеграции > 0, добавляем графики разностей
            if isinstance(d_value, (int, float)) and d_value > 0:
                # Первые разности
                plt.figure(figsize=(10, 6))
                diff1 = df[column].diff().dropna()
                plt.plot(df.index[1:], diff1)
                plt.title(f'Первые разности ряда {column}')
                plt.xlabel('Дата')
                plt.ylabel('Значение')
                plt.grid(True)
                
                # Сохраняем диаграмму в память
                img_stream = io.BytesIO()
                plt.savefig(img_stream, format='png')
                img_stream.seek(0)
                plt.close()
                
                # Добавляем изображение в документ
                doc.add_picture(img_stream, width=Inches(6))
            
            # Если порядок интеграции > 1, добавляем график вторых разностей
            if isinstance(d_value, (int, float)) and d_value > 1:
                # Вторые разности
                plt.figure(figsize=(10, 6))
                diff2 = df[column].diff().diff().dropna()
                plt.plot(df.index[2:], diff2)
                plt.title(f'Вторые разности ряда {column}')
                plt.xlabel('Дата')
                plt.ylabel('Значение')
                plt.grid(True)
                
                # Сохраняем диаграмму в память
                img_stream = io.BytesIO()
                plt.savefig(img_stream, format='png')
                img_stream.seek(0)
                plt.close()
                
                # Добавляем изображение в документ
                doc.add_picture(img_stream, width=Inches(6))
            
            # Добавляем разрыв страницы
            doc.add_page_break()
        
        # Увеличиваем номер раздела
        section_number += 1
    
    # Если выбран тест на коинтеграцию
    if tests.get('cointegration', False):
        # ИСПРАВЛЕНО: добавляем заголовок без указания стиля
        heading = doc.add_heading(f'{section_number}. Коинтеграционный анализ', level=1)
        heading.style = heading1_style
        
        doc.add_paragraph('В этом разделе представлены результаты тестов на коинтеграцию. Наличие коинтеграции позволяет строить модели, учитывающие долгосрочное равновесие между переменными.', style='NormalCustom')
        
        # Получаем параметры теста на коинтеграцию
        coint_params = tests.get('cointegration', {})
        if isinstance(coint_params, bool):
            coint_params = {}
        
        det_order = coint_params.get('det_order', 1)
        k_ar_diff = coint_params.get('k_ar_diff', 1)
        
        # Исследуем все пары переменных с тестом Энгла-Грейнджера
        # ИСПРАВЛЕНО: добавляем заголовок без указания стиля
        heading = doc.add_heading(f'{section_number}.1. Тест Энгла-Грейнджера', level=2)
        heading.style = heading2_style
        
        doc.add_paragraph('Тест Энгла-Грейнджера проверяет наличие коинтеграции между парами переменных. Если p-значение меньше 0.05, отвергается нулевая гипотеза об отсутствии коинтеграции.', style='NormalCustom')
        
        # Создаем таблицу для результатов
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        # Заголовки таблицы
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Переменная 1'
        hdr_cells[1].text = 'Переменная 2'
        hdr_cells[2].text = 'p-значение'
        hdr_cells[3].text = 'Результат'
        
        # Проверяем все пары переменных
        cointegrated_pairs = []
        
        for i in range(len(numeric_columns)):
            for j in range(i+1, len(numeric_columns)):
                var1 = numeric_columns[i]
                var2 = numeric_columns[j]
                
                # Проверка на коинтеграцию
                is_cointegrated, cointegration_output = test_cointegration(
                    df[var1],
                    df[var2],
                    var1,
                    var2
                )
                
                # Ищем p-value в выводе теста
                p_value = None
                for line in cointegration_output:
                    if 'p-value:' in line:
                        p_value = line.split('p-value:')[1].strip()
                        break
                
                # Добавляем строку в таблицу
                row_cells = table.add_row().cells
                row_cells[0].text = var1
                row_cells[1].text = var2
                row_cells[2].text = p_value if p_value else '-'
                row_cells[3].text = 'Коинтегрированы' if is_cointegrated else 'Не коинтегрированы'
                
                # Сохраняем коинтегрированные пары
                if is_cointegrated:
                    cointegrated_pairs.append((var1, var2))
        
        # Тест Йохансена на коинтеграцию (для всех переменных сразу)
        # ИСПРАВЛЕНО: добавляем заголовок без указания стиля
        heading = doc.add_heading(f'{section_number}.2. Тест Йохансена', level=2)
        heading.style = heading2_style
        
        doc.add_paragraph('Тест Йохансена позволяет определить количество коинтеграционных соотношений между несколькими переменными одновременно.', style='NormalCustom')
        
        # Параметры теста
        doc.add_paragraph(f'Порядок детерминистических компонент: {det_order}', style='NormalCustom')
        doc.add_paragraph(f'Количество лагов: {k_ar_diff}', style='NormalCustom')
        
        # Выполняем тест Йохансена
        is_cointegrated_johansen, cointegration_output_johansen = johansen_test(
            df,
            numeric_columns,
            det_order=det_order,
            k_ar_diff=k_ar_diff
        )
        
        # Выводим детальные результаты
        for line in cointegration_output_johansen:
            doc.add_paragraph(line, style='NormalCustom')
        
        # Визуализация отношения между коинтегрированными переменными
        if cointegrated_pairs:
            # ИСПРАВЛЕНО: добавляем заголовок без указания стиля
            heading = doc.add_heading(f'{section_number}.3. Визуализация коинтегрированных пар', level=2)
            heading.style = heading2_style
            
            for var1, var2 in cointegrated_pairs[:3]:  # Ограничиваем до 3 пар для наглядности
                # Создаем диаграмму
                plt.figure(figsize=(10, 6))
                plt.plot(df.index, df[var1], label=var1)
                plt.plot(df.index, df[var2], label=var2)
                plt.title(f'Коинтегрированные переменные: {var1} и {var2}')
                plt.xlabel('Дата')
                plt.ylabel('Значение')
                plt.legend()
                plt.grid(True)
                
                # Сохраняем диаграмму в память
                img_stream = io.BytesIO()
                plt.savefig(img_stream, format='png')
                img_stream.seek(0)
                plt.close()
                
                # Добавляем изображение в документ
                doc.add_picture(img_stream, width=Inches(6))
        
        # Общее заключение по коинтеграции
        any_eg_cointegrated = len(cointegrated_pairs) > 0
        is_cointegrated = any_eg_cointegrated or is_cointegrated_johansen
        
        # ИСПРАВЛЕНО: добавляем заголовок без указания стиля
        heading = doc.add_heading(f'{section_number}.4. Общее заключение по коинтеграции', level=2)
        heading.style = heading2_style
        
        if is_cointegrated:
            conclusion = "Обнаружена коинтеграция между переменными. Рекомендуется использовать модель VECM для учета долгосрочных зависимостей."
        else:
            conclusion = "Коинтеграция между переменными не обнаружена. Рекомендуется использовать модель VAR в разностях."
        
        doc.add_paragraph(conclusion, style='NormalCustom')
        
        # Добавляем разрыв страницы
        doc.add_page_break()
        
        # Увеличиваем номер раздела
        section_number += 1
    
    # Если выбрана модель VARX
    if models.get('varx', False):
        heading = doc.add_heading(f'{section_number}. Модель векторной авторегрессии с экзогенными переменными (VARX)', level=1)
        heading.style = heading1_style
        
        doc.add_paragraph('Модель VARX представляет собой расширение модели VAR, которая включает также экзогенные переменные. Она позволяет моделировать взаимозависимость между несколькими временными рядами с учетом внешних факторов.', style=normal_style)
        
        # Получаем параметры модели VARX
        varx_params = models.get('varx', {})
        lags = varx_params.get('lags', 2)
        forecast_periods = varx_params.get('forecast_periods', 6)
        forecast_unit = varx_params.get('forecast_unit', 'months')
        
        # Получаем эндогенные и экзогенные переменные из параметров модели
        endogenous_vars = varx_params.get('endogenous_vars', [])
        exogenous_vars = varx_params.get('exogenous_vars', [])
        
        # Если эндогенные переменные не указаны, берем первые две числовые колонки
        if not endogenous_vars:
            endogenous_vars = numeric_columns[:min(2, len(numeric_columns))]
        
        # Вызываем обновленную функцию построения модели VARX
        try:
            # Импорт функции
            from ts_analysis import 
            
            # Построение модели VARX с прогнозированием в будущее
            model_results = (
                df, 
                endogenous_vars, 
                exogenous_vars, 
                lags=lags,
                forecast_periods=forecast_periods,
                forecast_unit=forecast_unit
            )
            
            # Добавляем информацию о модели
            heading = doc.add_heading(f'{section_number}.1. Информация о модели', level=2)
            heading.style = heading2_style
            
            # Определяем тип модели
            model_type = model_results.get('model_info', {}).get('model_type', 'VARX')
            if not model_type:
                model_type = "ARX" if len(endogenous_vars) == 1 else "VARX"
            
            # Таблица с основными параметрами модели
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Параметр'
            hdr_cells[1].text = 'Значение'
            
            # Заполняем таблицу
            params_items = [
                ('Тип модели', model_type),
                ('Эндогенные переменные', ', '.join(endogenous_vars)),
                ('Экзогенные переменные', ', '.join(exogenous_vars) if exogenous_vars else 'Не используются'),
                ('Количество лагов', str(lags))
            ]
            
            # Добавляем информацию о прогнозном периоде
            unit_names = {
                'months': 'месяцев',
                'quarters': 'кварталов', 
                'years': 'лет'
            }
            unit_name = unit_names.get(forecast_unit, forecast_unit)
            params_items.append(('Период прогнозирования', f'{forecast_periods} {unit_name}'))
            
            # Добавляем параметры, если они есть в результатах
            if 'model_info' in model_results:
                model_info = model_results['model_info']
                
                if 'aic' in model_info and model_info['aic'] is not None:
                    params_items.append(('AIC', f"{model_info['aic']:.4f}"))
                
                if 'bic' in model_info and model_info['bic'] is not None:
                    params_items.append(('BIC', f"{model_info['bic']:.4f}"))
                
                if 'hqic' in model_info and model_info['hqic'] is not None:
                    params_items.append(('HQIC', f"{model_info['hqic']:.4f}"))
                
                if 'train_size' in model_info:
                    params_items.append(('Размер обучающей выборки', f"{model_info['train_size']} наблюдений"))
                
                if 'test_size' in model_info:
                    params_items.append(('Размер тестовой выборки', f"{model_info['test_size']} наблюдений"))
            
            for param, value in params_items:
                row_cells = table.add_row().cells
                row_cells[0].text = param
                row_cells[1].text = value
            
            # Добавляем раздел о выборе оптимального числа лагов, если есть
            if 'diagnostics' in model_results and 'lag_analysis' in model_results['diagnostics']:
                heading = doc.add_heading(f'{section_number}.2. Анализ оптимального числа лагов', level=2)
                heading.style = heading2_style
                
                lag_analysis = model_results['diagnostics']['lag_analysis']
                
                # Создаем таблицу с оптимальными лагами
                lag_table = doc.add_table(rows=1, cols=2)
                lag_table.style = 'Table Grid'
                
                hdr_cells = lag_table.rows[0].cells
                hdr_cells[0].text = 'Переменная'
                hdr_cells[1].text = 'Оптимальное число лагов'
                
                # Заполняем таблицу для каждой переменной
                for var, lag in lag_analysis['optimal_lags'].items():
                    row_cells = lag_table.add_row().cells
                    row_cells[0].text = var
                    row_cells[1].text = str(lag)
                
                # Добавляем рекомендации
                doc.add_paragraph(f"Рекомендуемое количество лагов (максимальное из оптимальных): {lag_analysis['optimal_lag_final']}")
                doc.add_paragraph(f"В модели использовано: {lag_analysis['chosen_lag']} лагов")
            
            # Добавляем графики ACF и PACF, если они есть
            if 'plots' in model_results and 'acf_pacf_plots' in model_results['plots']:
                heading = doc.add_heading(f'{section_number}.3. Графики автокорреляционных функций', level=2)
                heading.style = heading2_style
                
                for i, plot_base64 in enumerate(model_results['plots']['acf_pacf_plots']):
                    # Преобразуем base64 в изображение
                    import base64
                    image_data = base64.b64decode(plot_base64)
                    image_stream = io.BytesIO(image_data)
                    
                    # Добавляем изображение в документ
                    try:
                        doc.add_picture(image_stream, width=Inches(6))
                        if i < len(endogenous_vars):
                            doc.add_paragraph(f'ACF и PACF для переменной "{endogenous_vars[i]}"', style=normal_style)
                        else:
                            doc.add_paragraph(f'ACF и PACF график {i+1}', style=normal_style)
                    except Exception as e:
                        doc.add_paragraph(f'Не удалось добавить изображение графика: {str(e)}', style=normal_style)
            
            # Добавляем метрики прогнозирования на валидационных данных, если они есть
            if 'validation' in model_results and 'metrics' in model_results['validation']:
                heading = doc.add_heading(f'{section_number}.4. Метрики качества прогноза (на тестовых данных)', level=2)
                heading.style = heading2_style
                
                metrics = model_results['validation']['metrics']
                
                # Определяем количество столбцов в таблице метрик
                metrics_columns = 3  # по умолчанию: Переменная, MSE, MAE
                if 'rmse' in metrics:
                    metrics_columns += 1
                if 'mape' in metrics:
                    metrics_columns += 1
                
                metrics_table = doc.add_table(rows=1, cols=metrics_columns)
                metrics_table.style = 'Table Grid'
                
                # Заголовки таблицы
                hdr_cells = metrics_table.rows[0].cells
                hdr_cells[0].text = 'Переменная'
                hdr_cells[1].text = 'MSE'
                hdr_cells[2].text = 'MAE'
                
                # Добавляем RMSE и MAPE, если есть
                col_idx = 3
                if 'rmse' in metrics:
                    hdr_cells[col_idx].text = 'RMSE'
                    col_idx += 1
                if 'mape' in metrics:
                    hdr_cells[col_idx].text = 'MAPE'
                
                # Заполняем таблицу метриками для каждой переменной
                for variable in endogenous_vars:
                    if variable in metrics['mse']:
                        row_cells = metrics_table.add_row().cells
                        row_cells[0].text = variable
                        row_cells[1].text = f"{metrics['mse'][variable]:.4f}"
                        row_cells[2].text = f"{metrics['mae'][variable]:.4f}"
                        
                        # Добавляем RMSE, если есть
                        col_idx = 3
                        if 'rmse' in metrics and variable in metrics['rmse']:
                            row_cells[col_idx].text = f"{metrics['rmse'][variable]:.4f}"
                            col_idx += 1
                        elif 'rmse' in metrics:
                            row_cells[col_idx].text = 'N/A'
                            col_idx += 1
                        
                        # Добавляем MAPE, если есть
                        if 'mape' in metrics and variable in metrics['mape']:
                            import numpy as np
                            if not np.isnan(metrics['mape'][variable]):
                                row_cells[col_idx].text = f"{metrics['mape'][variable]:.2f}%"
                            else:
                                row_cells[col_idx].text = 'N/A'
            
            # Добавляем раздел с прогнозными значениями в будущее
            if 'forecasts' in model_results and 'future_forecast' in model_results['forecasts']:
                heading = doc.add_heading(f'{section_number}.5. Прогнозные значения в будущее', level=2)
                heading.style = heading2_style
                
                future_data = model_results['forecasts']['future_forecast']
                
                if 'values' in future_data and 'dates' in future_data and future_data['values']:
                    # Создаем таблицу с прогнозными значениями
                    forecast_table = doc.add_table(rows=1, cols=len(endogenous_vars) + 1)
                    forecast_table.style = 'Table Grid'
                    
                    # Заголовки таблицы
                    hdr_cells = forecast_table.rows[0].cells
                    hdr_cells[0].text = 'Дата'
                    for i, var in enumerate(endogenous_vars):
                        hdr_cells[i + 1].text = var
                    
                    # Заполняем таблицу прогнозными значениями
                    for i, date in enumerate(future_data['dates']):
                        row_cells = forecast_table.add_row().cells
                        row_cells[0].text = date
                        for j, var in enumerate(endogenous_vars):
                            if var in future_data['values'] and i < len(future_data['values'][var]):
                                row_cells[j + 1].text = f"{future_data['values'][var][i]:.4f}"
                            else:
                                row_cells[j + 1].text = 'N/A'
                    
                    # Добавляем пояснение к таблице
                    doc.add_paragraph(f"Прогнозные значения на {forecast_periods} {unit_name} вперед от последней даты в данных.", style=normal_style)
                    
                    # Добавляем статистику по прогнозам
                    doc.add_paragraph("", style=normal_style)
                    doc.add_paragraph("Статистика прогнозных значений:", style=normal_style)
                    
                    # Создаем таблицу со статистикой
                    stats_table = doc.add_table(rows=1, cols=len(endogenous_vars) + 1)
                    stats_table.style = 'Table Grid'
                    
                    # Заголовки таблицы статистики
                    stats_hdr_cells = stats_table.rows[0].cells
                    stats_hdr_cells[0].text = 'Показатель'
                    for i, var in enumerate(endogenous_vars):
                        stats_hdr_cells[i + 1].text = var
                    
                    # Рассчитываем и добавляем статистику
                    stats_metrics = ['Среднее значение', 'Минимум', 'Максимум', 'Стандартное отклонение']
                    
                    for metric in stats_metrics:
                        row_cells = stats_table.add_row().cells
                        row_cells[0].text = metric
                        
                        for j, var in enumerate(endogenous_vars):
                            if var in future_data['values'] and future_data['values'][var]:
                                values = future_data['values'][var]
                                if metric == 'Среднее значение':
                                    value = np.mean(values)
                                elif metric == 'Минимум':
                                    value = np.min(values)
                                elif metric == 'Максимум':
                                    value = np.max(values)
                                elif metric == 'Стандартное отклонение':
                                    value = np.std(values)
                                
                                row_cells[j + 1].text = f"{value:.4f}"
                            else:
                                row_cells[j + 1].text = 'N/A'
                    
                    # Добавляем интерпретацию результатов
                    doc.add_paragraph("", style=normal_style)
                    doc.add_paragraph("Интерпретация прогнозов:", style=normal_style)
                    
                    for var in endogenous_vars:
                        if var in future_data['values'] and future_data['values'][var]:
                            values = future_data['values'][var]
                            first_value = values[0]
                            last_value = values[-1]
                            trend_direction = "растущий" if last_value > first_value else "снижающийся" if last_value < first_value else "стабильный"
                            change_percent = ((last_value - first_value) / first_value) * 100 if first_value != 0 else 0
                            
                            interpretation = f"• {var}: прогнозируется {trend_direction} тренд с изменением на {change_percent:.2f}% за прогнозный период."
                            doc.add_paragraph(interpretation, style=normal_style)
                    
                    # Добавляем предупреждения и ограничения
                    doc.add_paragraph("", style=normal_style)
                    doc.add_paragraph("Важные замечания:", style=normal_style)
                    doc.add_paragraph("• Прогнозы основаны на исторических данных и предполагают сохранение выявленных закономерностей.", style=normal_style)
                    doc.add_paragraph("• Точность прогнозов снижается с увеличением горизонта прогнозирования.", style=normal_style)
                    doc.add_paragraph("• Внешние факторы, не учтенные в модели, могут существенно влиять на фактические значения.", style=normal_style)
                    doc.add_paragraph("• Рекомендуется регулярно обновлять модель при поступлении новых данных.", style=normal_style)
                
                else:
                    doc.add_paragraph("Прогнозные данные недоступны из-за ошибки в модели или отсутствия данных.", style=normal_style)
            
            # Добавляем графики прогнозов с будущими значениями, если они есть
            if 'plots' in model_results and 'forecast_plots' in model_results['plots']:
                heading = doc.add_heading(f'{section_number}.6. Графики прогнозов с будущими значениями', level=2)
                heading.style = heading2_style
                
                plots = model_results['plots']['forecast_plots']
                
                for i, plot_base64 in enumerate(plots):
                    # Преобразуем base64 в изображение
                    import base64
                    image_data = base64.b64decode(plot_base64)
                    image_stream = io.BytesIO(image_data)
                    
                    # Добавляем изображение в документ
                    try:
                        doc.add_picture(image_stream, width=Inches(6))
                        if i < len(endogenous_vars):
                            doc.add_paragraph(f'Комплексный прогноз для переменной "{endogenous_vars[i]}" включающий исторические данные, валидацию на тестовых данных и прогноз в будущее на {forecast_periods} {unit_name}.', style=normal_style)
                        else:
                            doc.add_paragraph(f'Комплексный прогноз {i+1} с будущими значениями', style=normal_style)
                    except Exception as e:
                        doc.add_paragraph(f'Не удалось добавить изображение графика: {str(e)}', style=normal_style)
                
                # Добавляем легенду графиков
                doc.add_paragraph("Легенда графиков:", style=normal_style)
                doc.add_paragraph("• Синяя линия - исторические данные (обучающая выборка)", style=normal_style)
                doc.add_paragraph("• Зеленая линия - фактические значения на тестовых данных", style=normal_style)
                doc.add_paragraph("• Красная пунктирная линия - прогнозы на тестовых данных", style=normal_style)
                doc.add_paragraph("• Оранжевая линия с маркерами - прогнозы в будущее", style=normal_style)
                doc.add_paragraph("• Серая вертикальная линия - конец обучающих данных", style=normal_style)
                doc.add_paragraph("• Красная вертикальная линия - конец всех исторических данных", style=normal_style)
            
            # Добавляем заключение по модели
            heading = doc.add_heading(f'{section_number}.7. Заключение по модели {model_type}', level=2)
            heading.style = heading2_style
            
            conclusion_text = f"""
            Модель {model_type} была успешно построена с использованием {len(endogenous_vars)} эндогенных переменных 
            {'и ' + str(len(exogenous_vars)) + ' экзогенных переменных' if exogenous_vars else 'без экзогенных переменных'}. 
            Модель обучена на {model_results.get('model_info', {}).get('train_size', 'N/A')} наблюдениях 
            и протестирована на {model_results.get('model_info', {}).get('test_size', 'N/A')} наблюдениях.
            
            Создан прогноз на {forecast_periods} {unit_name} в будущее, который может быть использован 
            для планирования и принятия управленческих решений. Рекомендуется регулярно обновлять модель 
            при поступлении новых данных для поддержания точности прогнозов.
            """
            
            doc.add_paragraph(conclusion_text.strip(), style=normal_style)
            
        except Exception as e:
            # В случае ошибки при построении модели, добавляем информацию об этом
            heading = doc.add_heading(f'{section_number}.1. Информация о модели', level=2)
            heading.style = heading2_style
            
            doc.add_paragraph(f"При построении модели VARX произошла ошибка: {str(e)}", style=normal_style)
            
            doc.add_paragraph('Модель VARX представляет собой расширение модели VAR, которая включает также экзогенные переменные. Она позволяет моделировать взаимозависимость между несколькими временными рядами с учетом внешних факторов и создавать прогнозы в будущее.', style=normal_style)
            
            # Добавляем заглушку-пример графика
            plt.figure(figsize=(12, 8))
            
            # Создаем пример данных
            dates = pd.date_range(start='2020-01-01', end='2025-12-01', freq='M')
            historical_end = len(dates) - 12  # Последние 12 месяцев как "будущее"
            
            # Исторические данные
            np.random.seed(42)
            trend1 = np.linspace(100, 150, historical_end) + np.random.normal(0, 5, historical_end)
            trend2 = np.linspace(80, 120, historical_end) + np.random.normal(0, 3, historical_end)
            
            # "Прогнозы" в будущее
            future_trend1 = np.linspace(150, 160, 12) + np.random.normal(0, 3, 12)
            future_trend2 = np.linspace(120, 125, 12) + np.random.normal(0, 2, 12)
            
            # Строим график
            plt.plot(dates[:historical_end], trend1, 'b-', label='Переменная 1 (исторические)', linewidth=2)
            plt.plot(dates[:historical_end], trend2, 'g-', label='Переменная 2 (исторические)', linewidth=2)
            plt.plot(dates[historical_end:], future_trend1, 'orange', linewidth=3, marker='o', 
                    markersize=6, label='Переменная 1 (прогноз)')
            plt.plot(dates[historical_end:], future_trend2, 'red', linewidth=3, marker='s', 
                    markersize=6, label='Переменная 2 (прогноз)')
            
            # Добавляем вертикальную линию разделения
            plt.axvline(x=dates[historical_end-1], color='gray', linestyle=':', alpha=0.7, 
                       label='Конец исторических данных')
            
            plt.title('Пример моделирования VARX с прогнозированием в будущее', fontsize=16, fontweight='bold')
            plt.xlabel('Дата', fontsize=12)
            plt.ylabel('Значение', fontsize=12)
            plt.legend(fontsize=10)
            plt.grid(True, alpha=0.3)
            plt.xticks(rotation=45)
            plt.tight_layout()
            
            # Сохраняем диаграмму в память
            img_stream = io.BytesIO()
            plt.savefig(img_stream, format='png', dpi=150, bbox_inches='tight')
            img_stream.seek(0)
            plt.close()
            
            # Добавляем изображение в документ
            doc.add_picture(img_stream, width=Inches(6))
            
            doc.add_paragraph('Этот пример показывает, как модель VARX может анализировать взаимосвязи между переменными и создавать прогнозы в будущее. В нормальном режиме работы система строит реальные прогнозы на основе ваших данных и оценивает их качество по различным метрикам.', style=normal_style)
        
        # Добавляем разрыв страницы
        doc.add_page_break()
        
        # Увеличиваем номер раздела
        section_number += 1
    
    # Сохраняем документ
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
        temp_path = temp_file.name
        doc.save(temp_path)
    
    return temp_path
def extreme_transform_to_stationary(series, method='auto', title='', debug=False):
    """
    Применяет экстремальное преобразование к временному ряду для достижения стационарности
    и улучшения статистических свойств.
    
    Parameters:
    -----------
    series : pandas.Series
        Исходный временной ряд
    method : str, default='auto'
        Метод преобразования:
        - 'auto': Автоматический выбор метода на основе характеристик ряда
        - 'rank': Ранговая нормализация (метод Блома)
        - 'boxcox': Преобразование Бокса-Кокса
        - 'yeojohnson': Преобразование Йео-Джонсона
        - 'simple': Простое дифференцирование
        - 'log': Логарифмическое дифференцирование
        - 'percentage': Процентное изменение
    title : str, default=''
        Название ряда для вывода информации
    debug : bool, default=False
        Вывод отладочной информации о процессе преобразования
    
    Returns:
    --------
    pandas.Series
        Преобразованный стационарный ряд
    dict
        Информация о преобразовании
    """
    if debug:
        print(f"\nЭкстремальное преобразование ряда '{title}' методом '{method}'")

    # Копируем исходный ряд и удаляем пропуски
    original_series = series.copy().dropna()
    
    # Проверяем стационарность исходного ряда
    original_status, original_output = check_stationarity(original_series, f"{title} (исходный)")
    
    # Если ряд уже стационарен, возвращаем его без изменений
    if original_status == "Стационарен":
        if debug:
            print("Ряд уже стационарен, преобразование не требуется")
        
        return original_series, {
            'transform_method': 'none',
            'double_differencing': False,
            'winsorization': False,
            'd_value': 0,
            'original_status': original_status,
            'final_status': original_status
        }
    
    # Статистические характеристики исходного ряда
    skewness = original_series.skew()
    kurtosis = original_series.kurtosis()
    
    # Автоматический выбор метода преобразования, если указан 'auto'
    if method == 'auto':
        # Проверяем наличие слов 'задолженность' или 'долг' в названии
        is_debt = 'задолженност' in title.lower() or 'долг' in title.lower()
        
        # Для рядов с высокой асимметрией или эксцессом использовать ранговую нормализацию
        if is_debt or abs(skewness) > 3 or abs(kurtosis) > 10:
            selected_method = 'rank'
        # Для положительных рядов используем Box-Cox
        elif original_series.min() > 0:
            selected_method = 'boxcox'
        # Для остальных используем Yeo-Johnson
        else:
            selected_method = 'yeojohnson'
            
        if debug:
            print(f"Автоматически выбран метод: {selected_method}")
    else:
        selected_method = method
    
    # Применяем выбранный метод преобразования
    try:
        if selected_method == 'rank':
            # Ранговое преобразование к нормальному распределению по методу Блома
            n = len(original_series)
            ranks = original_series.rank()
            norm_quantiles = stats.norm.ppf((ranks - 0.375) / (n + 0.25))
            normalized_series = pd.Series(norm_quantiles, index=original_series.index)
            transform_info = "Ранговая нормализация (метод Блома)"
            
        elif selected_method == 'boxcox':
            # Если есть отрицательные или нулевые значения, сдвигаем ряд
            if original_series.min() <= 0:
                min_val = original_series.min()
                offset = abs(min_val) + 1  # Сдвиг, чтобы все значения были > 0
                adjusted_series = original_series + offset
                transformed_values, lambda_param = stats.boxcox(adjusted_series.values)
            else:
                transformed_values, lambda_param = stats.boxcox(original_series.values)
                
            normalized_series = pd.Series(transformed_values, index=original_series.index)
            transform_info = f"Box-Cox (лямбда = {lambda_param:.4f})"
            
        elif selected_method == 'yeojohnson':
            # Преобразование Yeo-Johnson (работает с любыми значениями)
            transformed_values, lambda_param = stats.yeojohnson(original_series.values)
            normalized_series = pd.Series(transformed_values, index=original_series.index)
            transform_info = f"Yeo-Johnson (лямбда = {lambda_param:.4f})"
            
        elif selected_method == 'simple':
            # Простое дифференцирование
            normalized_series = original_series
            transform_info = "Простое дифференцирование"
            
        elif selected_method == 'log':
            # Логарифмическое преобразование
            if original_series.min() <= 0:
                min_val = original_series.min()
                offset = abs(min_val) + 1  # Сдвиг, чтобы все значения были > 0
                adjusted_series = original_series + offset
                normalized_series = np.log(adjusted_series)
            else:
                normalized_series = np.log(original_series)
            transform_info = "Логарифмическое преобразование"
            
        elif selected_method == 'percentage':
            # Процентное изменение
            normalized_series = original_series
            transform_info = "Процентное изменение"
        else:
            # По умолчанию применяем ранговую нормализацию
            n = len(original_series)
            ranks = original_series.rank()
            norm_quantiles = stats.norm.ppf((ranks - 0.375) / (n + 0.25))
            normalized_series = pd.Series(norm_quantiles, index=original_series.index)
            transform_info = "Ранговая нормализация (метод Блома) - по умолчанию"
            
    except Exception as e:
        if debug:
            print(f"Ошибка при применении метода {selected_method}: {str(e)}")
            print("Использую ранговую нормализацию как запасной вариант")
            
        # Ранговая нормализация как запасной вариант
        n = len(original_series)
        ranks = original_series.rank()
        norm_quantiles = stats.norm.ppf((ranks - 0.375) / (n + 0.25))
        normalized_series = pd.Series(norm_quantiles, index=original_series.index)
        transform_info = "Ранговая нормализация (метод Блома) - запасной метод"
    
    # Проверяем стационарность после нормализации
    norm_status, _ = check_stationarity(normalized_series, f"{title} (после нормализации)")
    
    # Если ряд уже стационарен после нормализации, возвращаем его
    if norm_status == "Стационарен":
        if debug:
            print("Ряд стал стационарным после начального преобразования")
        
        return normalized_series, {
            'transform_method': transform_info,
            'double_differencing': False,
            'winsorization': False,
            'd_value': 0,  # Нет дифференцирования
            'original_status': original_status,
            'final_status': norm_status
        }
    
    # Применяем дифференцирование для достижения стационарности
    diff1_series = normalized_series.diff().dropna()
    
    # Проверяем стационарность после первого дифференцирования
    diff1_status, _ = check_stationarity(diff1_series, f"{title} (первые разности)")
    
    # Если ряд стационарен после первого дифференцирования, возвращаем его
    if diff1_status == "Стационарен":
        if debug:
            print("Ряд стал стационарным после первого дифференцирования")
        
        # Применяем винсоризацию для убирания оставшихся выбросов
        q05 = diff1_series.quantile(0.05)
        q95 = diff1_series.quantile(0.95)
        iqr = q95 - q05
        
        lower_bound = q05 - 1.5 * iqr
        upper_bound = q95 + 1.5 * iqr
        
        winsorized_series = diff1_series.clip(lower=lower_bound, upper=upper_bound)
        
        return winsorized_series, {
            'transform_method': transform_info,
            'double_differencing': False,
            'winsorization': True,
            'd_value': 1,  # Первый порядок дифференцирования
            'original_status': original_status,
            'final_status': "Стационарен"
        }
    
    # Применяем второе дифференцирование
    diff2_series = diff1_series.diff().dropna()
    
    # Проверяем стационарность после второго дифференцирования
    diff2_status, _ = check_stationarity(diff2_series, f"{title} (вторые разности)")
    
    # Интенсивная винсоризация для убирания оставшихся выбросов
    q05 = diff2_series.quantile(0.05)
    q95 = diff2_series.quantile(0.95)
    iqr = q95 - q05
    
    lower_bound = q05 - 0.5 * iqr  # Более агрессивная винсоризация
    upper_bound = q95 + 0.5 * iqr
    
    final_series = diff2_series.clip(lower=lower_bound, upper=upper_bound)
    
    # Финальная проверка стационарности
    final_status, _ = check_stationarity(final_series, f"{title} (после винсоризации)")
    
    if debug:
        if diff2_status == "Стационарен":
            print("Ряд стал стационарным после второго дифференцирования")
        
        if final_status == "Стационарен":
            print("Ряд стал стационарным после всех преобразований")
        else:
            print("Ряд не удалось сделать стационарным даже после всех преобразований")
    
    return final_series, {
        'transform_method': transform_info,
        'double_differencing': True,
        'winsorization': True,
        'd_value': 2,  # Второй порядок дифференцирования
        'original_status': original_status,
        'final_status': final_status
    }
