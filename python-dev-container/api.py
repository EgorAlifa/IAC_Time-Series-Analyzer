from fastapi import FastAPI, UploadFile, File, HTTPException, Form, Request
from fastapi.responses import JSONResponse, FileResponse, StreamingResponse
from pydantic import BaseModel
from typing import List, Optional, Dict, Any, Union
from fastapi import Query
from datetime import datetime, timedelta
import pandas as pd
import numpy as np
import tempfile
import os
import json
import requests
from io import BytesIO
import uuid
import shutil
import threading
import time
from pathlib import Path

from ts_analysis import analyze_data, create_word_report, test_cointegration, johansen_test, save_message, cleanup_old_messages, FeedbackForm, transform_to_stationary, transform_to_first_order, transform_to_second_order, format_time_series_for_preview, extreme_transform_to_stationary, build_varx_model_with_future_forecast

app = FastAPI(title="Анализ стационарности временных рядов")
from fastapi.middleware.cors import CORSMiddleware

@app.middleware("http")
async def add_cors_headers(request: Request, call_next):
    if request.method == "OPTIONS":
        # Обработка preflight запросов
        response = Response()
        response.headers["Access-Control-Allow-Origin"] = "*"
        response.headers["Access-Control-Allow-Methods"] = "GET, POST, PUT, DELETE, OPTIONS"
        response.headers["Access-Control-Allow-Headers"] = "*"
        response.headers["Access-Control-Max-Age"] = "3600"
        return response
    
    try:
        response = await call_next(request)
    except Exception as e:
        # Даже при ошибке отправляем CORS заголовки
        response = JSONResponse(
            status_code=500,
            content={"error": f"Internal server error: {str(e)}"}
        )
    
    # Добавляем CORS заголовки ко всем ответам
    response.headers["Access-Control-Allow-Origin"] = "*"
    response.headers["Access-Control-Allow-Methods"] = "GET, POST, PUT, DELETE, OPTIONS"
    response.headers["Access-Control-Allow-Headers"] = "*"
    response.headers["Access-Control-Allow-Credentials"] = "true"
    
    return response

# ТАКЖЕ оставьте стандартный CORS middleware:
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Создаем директорию для временного хранения файлов, если ее нет
TEMP_FILES_DIR = Path("temp_files")
TEMP_FILES_DIR.mkdir(exist_ok=True)

# Словарь для хранения времени загрузки файлов
file_timestamps = {}

# Функция для удаления просроченных файлов
def cleanup_old_files():
    while True:
        current_time = time.time()
        expired_files = []

        # Находим все файлы старше 1 часа
        for file_id, timestamp in list(file_timestamps.items()):
            if current_time - timestamp > 3600:  # 3600 секунд = 1 час
                expired_files.append(file_id)
        
        # Удаляем просроченные файлы
        for file_id in expired_files:
            file_path = TEMP_FILES_DIR / file_id
            if file_path.exists():
                os.remove(file_path)
            file_timestamps.pop(file_id, None)
            print(f"Файл {file_id} удален после истечения срока хранения")
        
        # Проверяем каждые 5 минут
        time.sleep(300)

# Запускаем поток для очистки файлов
cleanup_thread = threading.Thread(target=cleanup_old_files, daemon=True)
cleanup_thread.start()

class AnalysisResult(BaseModel):
    variable_results: Dict[str, Dict[str, Any]]
    cointegration_results: Optional[Dict[str, Any]] = None
def unpack_variables(variables):
    """
    Распаковывает переменные из возможных форматов (строка, список, JSON-строка)
    
    Args:
        variables: Переменные в различных форматах
        
    Returns:
        list: Список переменных
    """
    # Если None, возвращаем пустой список
    if variables is None:
        return []
    
    # Если строка, проверяем, является ли она JSON
    if isinstance(variables, str):
        try:
            # Пробуем распарсить как JSON
            import json
            parsed = json.loads(variables)
            if isinstance(parsed, list):
                return parsed
            else:
                return [variables]  # Не список JSON, возвращаем как есть
        except:
            # Не JSON, возвращаем как есть
            return [variables]
    
    # Если список, проверяем каждый элемент на JSON
    if isinstance(variables, list):
        result = []
        for item in variables:
            if isinstance(item, str):
                try:
                    # Пробуем распарсить как JSON
                    import json
                    parsed = json.loads(item)
                    if isinstance(parsed, list):
                        result.extend(parsed)  # Добавляем все элементы из JSON-массива
                    else:
                        result.append(parsed)  # Добавляем одиночное значение
                except:
                    # Не JSON, добавляем как есть
                    result.append(item)
            else:
                # Не строка, добавляем как есть
                result.append(item)
        return result
    
    # Для других типов возвращаем в списке
    return [variables]
@app.get("/api")
async def root():
    return {"message": "API для анализа стационарности временных рядов"}

@app.get("/api/news")
async def get_news_endpoint(
    q: str = Query(default="Ставка", description="Поисковый запрос"),
    limit: int = Query(default=7, description="Лимит количества новостей")
):
    try:
        print(f"Запрос новостей с параметрами: q={q}, limit={limit}")
        
        # API ключ для NewsAPI
        api_key = "1709ee5d87c94b238a395fefc8375b92"
        
        # Вычисляем дату 20 дней назад
        current_date = datetime.now()
        previous_date = current_date - timedelta(days=20)
        formatted_date = previous_date.strftime('%Y-%m-%d')
        
        # Формируем URL для запроса к NewsAPI
        url = f'https://newsapi.org/v2/everything?q={q}&from={formatted_date}&sortBy=publishedAt&apiKey={api_key}'
        print(f"URL запроса: {url}")
        
        # Выполняем запрос только один раз
        response = requests.get(url)
        print(f"Статус ответа: {response.status_code}")
        
        data = response.json()
        print(f"Получены данные: {data.get('status')}")
        
        news_list = []
        if data.get('status') == 'ok' and 'articles' in data:
            articles = data['articles']
            # Фильтруем статьи
            filtered_articles = [
                article for article in articles
                if "путин" not in article.get('title', '').lower() and "война" not in article.get('title', '').lower() and "войны" not in article.get('title', '').lower()
            ]
            
            # Ограничиваем количество новостей
            filtered_articles = filtered_articles[:limit]
            
            for article in filtered_articles:
                news_list.append({
                    'title': article.get('title', 'Без заголовка'),
                    'description': article.get('description', ''),
                    'url': article.get('url', '#'),
                    'publishedAt': article.get('publishedAt', ''),
                    'source': article.get('source', {}).get('name', 'Неизвестный источник')
                })
        else:
            # Возвращаем информацию об ошибке
            return {
                'error': 'Failed to fetch news',
                'details': data.get('message', 'Unknown error'),
                'status': data.get('status', 'error')
            }
        
        return news_list
    
    except Exception as e:
        print(f"Ошибка при получении новостей: {str(e)}")
        import traceback
        traceback.print_exc()
        return {
            'error': 'Exception occurred',
            'message': str(e)
        }

@app.post("/cointegration-analysis")
async def cointegration_analysis(
    file_id: str = Form(...),
    date_column: str = Form("Дата"),
    endogenous_vars: Optional[Union[List[str], str]] = Form(None),
    det_order: int = Form(1),  # Значение по умолчанию 1
    k_ar_diff: int = Form(1)   # Значение по умолчанию 1
):
    """
    Анализ коинтеграции для выбранных переменных
    
    - **file_id**: Идентификатор файла, полученный при загрузке
    - **date_column**: Имя столбца с датами (по умолчанию "Дата")
    - **endogenous_vars**: Список имен эндогенных переменных для анализа коинтеграции
    """
    # Распаковываем переменные из возможных форматов
    unpacked_vars = unpack_variables(endogenous_vars)
    
    # Добавим отладочный вывод
    print(f"Cointegration analysis - file_id: {file_id}")
    print(f"Cointegration analysis - date_column: {date_column}")
    print(f"Cointegration analysis - endogenous_vars: {endogenous_vars}")
    print(f"Cointegration analysis - unpacked_vars: {unpacked_vars}")
    
    # Проверяем наличие хотя бы двух переменных
    if len(unpacked_vars) < 2:
        return JSONResponse(
            status_code=400,
            content={"error": "Для анализа коинтеграции необходимо минимум две переменные"}
        )
    
    file_path = TEMP_FILES_DIR / file_id
    
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="Файл не найден или срок его хранения истек")
    
    try:
        # Определяем тип файла по расширению
        file_extension = os.path.splitext(str(file_path))[1].lower()
        
        # Чтение данных из файла
        if file_extension == '.csv':
            df = pd.read_csv(file_path)
        elif file_extension in ['.xlsx', '.xls']:
            df = pd.read_excel(file_path)
        else:
            # Пытаемся определить тип файла автоматически
            try:
                df = pd.read_excel(file_path)
            except:
                try:
                    df = pd.read_csv(file_path)
                except:
                    raise HTTPException(status_code=400, detail="Неподдерживаемый формат файла")
        
        # Проверка наличия столбца с датами
        if date_column not in df.columns:
            return JSONResponse(
                status_code=400,
                content={"error": f"Столбец {date_column} не найден в файле", "columns": df.columns.tolist()}
            )
            
        # Преобразование столбца даты
        try:
            df[date_column] = pd.to_datetime(df[date_column])
            df.set_index(date_column, inplace=True)
        except Exception as e:
            return JSONResponse(
                status_code=400,
                content={"error": f"Ошибка при преобразовании столбца даты: {str(e)}"}
            )
        
        # Проверяем, что все указанные переменные есть в данных
        missing_vars = [var for var in unpacked_vars if var not in df.columns]
        if missing_vars:
            return JSONResponse(
                status_code=400,
                content={
                    "error": f"Следующие переменные не найдены в данных: {', '.join(missing_vars)}",
                    "available_columns": df.columns.tolist()
                }
            )
        
        # Выполняем только коинтеграционный анализ
        
        # Для анализа методом Энгла-Грейнджера создаем все возможные пары
        eg_results = []
        
        for i in range(len(unpacked_vars)):
            for j in range(i+1, len(unpacked_vars)):
                var1 = unpacked_vars[i]
                var2 = unpacked_vars[j]
                
                is_cointegrated_eg, cointegration_output_eg = test_cointegration(
                    df[var1],
                    df[var2],
                    var1,
                    var2
                )
                
                eg_results.append({
                    'variables': [var1, var2],
                    'is_cointegrated': is_cointegrated_eg,
                    'details': cointegration_output_eg
                })
        
        # Тест Йохансена на коинтеграцию (для всех переменных сразу)
        is_cointegrated_johansen, cointegration_output_johansen = johansen_test(
            df,
            unpacked_vars,
            det_order=det_order,  # Используем переданное значение
            k_ar_diff=k_ar_diff   # Используем переданное значение
        )
        
        # Определяем общий результат коинтеграции
        any_eg_cointegrated = any([result['is_cointegrated'] for result in eg_results])
        is_cointegrated = any_eg_cointegrated or is_cointegrated_johansen
        
        # Формируем результат
        result = {
            'engle_granger_pairs': eg_results,
            'engle_granger_summary': {
                'is_cointegrated': any_eg_cointegrated,
                'cointegrated_pairs': [f"{result['variables'][0]} и {result['variables'][1]}" 
                                     for result in eg_results if result['is_cointegrated']]
            },
            'johansen': {
                'is_cointegrated': is_cointegrated_johansen,
                'details': cointegration_output_johansen
            },
            'conclusion': "Переменные коинтегрированы." if is_cointegrated else "Переменные не коинтегрированы.",
            'recommended_model': "VECM" if is_cointegrated else "VAR в разностях"
        }
        
        return result
        
    except Exception as e:
        import traceback
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Ошибка при анализе коинтеграции: {str(e)}")
@app.post("/get-time-series-data")
async def get_time_series_data(
    file_id: str = Form(...),
    date_column: str = Form("Дата"),
    column: str = Form(...)
):
    """
    Получение данных временного ряда для построения графиков
    """
    file_path = TEMP_FILES_DIR / file_id

    if not file_path.exists():
        raise HTTPException(status_code=404, detail="Файл не найден")

    try:
        # Чтение данных из файла
        if file_path.suffix.lower() == '.csv':
            df = pd.read_csv(file_path)
        else:
            df = pd.read_excel(file_path)

        # Проверка наличия столбцов
        if date_column not in df.columns or column not in df.columns:
            return JSONResponse(
                status_code=400,
                content={"error": f"Столбцы не найдены"}
            )

        # Преобразование даты и сортировка
        df[date_column] = pd.to_datetime(df[date_column])
        df = df.sort_values(by=date_column)

        # Формирование данных временного ряда
        time_series_data = []
        for _, row in df.iterrows():
            time_series_data.append({
                "date": row[date_column].strftime("%Y-%m-%d"),
                "value": float(row[column])
            })

        return time_series_data

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Ошибка при получении данных: {str(e)}")
@app.post("/upload-temp-file")
async def upload_temp_file(file: UploadFile = File(...)):
    """
    Загрузка файла для временного хранения (1 час)
    
    - **file**: Excel или CSV файл с данными
    
    Возвращает уникальный идентификатор файла для последующего доступа
    """
    try:
        # Генерируем уникальный ID для файла
        file_id = str(uuid.uuid4())
        file_path = TEMP_FILES_DIR / file_id
        
        # Сохраняем файл
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        
        # Запоминаем время загрузки
        file_timestamps[file_id] = time.time()
        
        # Определяем тип файла
        content_type = "unknown"
        file_extension = os.path.splitext(file.filename)[1].lower()
        if file_extension == '.csv':
            content_type = "text/csv"
        elif file_extension in ['.xlsx', '.xls']:
            content_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        
        return {
            "file_id": file_id,
            "original_filename": file.filename,
            "content_type": content_type,
            "expiration": "1 hour"
        }
    
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Ошибка при загрузке файла: {str(e)}")

@app.get("/temp-file/{file_id}")
async def get_temp_file(file_id: str):
    """
    Получение временного файла по его идентификатору
    
    - **file_id**: Идентификатор файла, полученный при загрузке
    """
    file_path = TEMP_FILES_DIR / file_id
    
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="Файл не найден или срок его хранения истек")
    
    return FileResponse(path=str(file_path))

@app.post("/analyze", response_model=AnalysisResult)
async def analyze_file(
    file: UploadFile = File(...),
    date_column: str = Form("Дата"),
    endogenous_vars: Optional[List[str]] = Form(None)
):
    """
    Анализ стационарности и коинтеграции временных рядов из Excel или CSV файла
    
    - **file**: Excel или CSV файл с данными
    - **date_column**: Имя столбца с датами (по умолчанию "Дата")
    - **endogenous_vars**: Список имен эндогенных переменных для анализа коинтеграции
    """
    # Проверка расширения файла
    file_extension = os.path.splitext(file.filename)[1].lower()
    
    try:
        # Чтение данных из файла
        content = await file.read()
        if file_extension == '.csv':
            df = pd.read_csv(BytesIO(content))
        elif file_extension in ['.xlsx', '.xls']:
            df = pd.read_excel(BytesIO(content))
        else:
            raise HTTPException(status_code=400, detail="Поддерживаются только файлы CSV, XLS или XLSX")
            
        # Проверка наличия столбца с датами
        if date_column not in df.columns:
            return JSONResponse(
                status_code=400,
                content={"error": f"Столбец {date_column} не найден в файле", "columns": df.columns.tolist()}
            )
            
        # Преобразование столбца даты
        try:
            df[date_column] = pd.to_datetime(df[date_column])
            df.set_index(date_column, inplace=True)
        except Exception as e:
            return JSONResponse(
                status_code=400,
                content={"error": f"Ошибка при преобразовании столбца даты: {str(e)}"}
            )
            
        # Проведение анализа
        results = analyze_data(df, endogenous_vars)
        
        return results
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Ошибка при анализе данных: {str(e)}")

@app.post("/analyze-with-file-id", response_model=AnalysisResult)
async def analyze_with_file_id(
    file_id: str = Form(...),
    date_column: str = Form("Дата"),
    endogenous_vars: Optional[Union[List[str], str]] = Form(None)
):
    """
    Анализ временных рядов из ранее загруженного файла
    
    - **file_id**: Идентификатор файла, полученный при загрузке
    - **date_column**: Имя столбца с датами (по умолчанию "Дата")
    - **endogenous_vars**: Список имен эндогенных переменных для анализа коинтеграции
    """
    # Распаковываем переменные из возможных форматов
    unpacked_vars = unpack_variables(endogenous_vars)
    
    # Добавим отладочный вывод
    print(f"Received file_id: {file_id}")
    print(f"Received date_column: {date_column}")
    print(f"Received endogenous_vars: {endogenous_vars}")
    print(f"Type of endogenous_vars: {type(endogenous_vars)}")
    print(f"Unpacked vars: {unpacked_vars}")
    
    file_path = TEMP_FILES_DIR / file_id
    
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="Файл не найден или срок его хранения истек")
    
    try:
        # Определяем тип файла по расширению
        file_extension = os.path.splitext(str(file_path))[1].lower()
        
        # Чтение данных из файла
        if file_extension == '.csv':
            df = pd.read_csv(file_path)
        elif file_extension in ['.xlsx', '.xls']:
            df = pd.read_excel(file_path)
        else:
            # Пытаемся определить тип файла автоматически
            try:
                df = pd.read_excel(file_path)
            except:
                try:
                    df = pd.read_csv(file_path)
                except:
                    raise HTTPException(status_code=400, detail="Неподдерживаемый формат файла")
        
        # Проверка наличия столбца с датами
        if date_column not in df.columns:
            return JSONResponse(
                status_code=400,
                content={"error": f"Столбец {date_column} не найден в файле", "columns": df.columns.tolist()}
            )
            
        # Преобразование столбца даты
        try:
            df[date_column] = pd.to_datetime(df[date_column])
            df.set_index(date_column, inplace=True)
        except Exception as e:
            return JSONResponse(
                status_code=400,
                content={"error": f"Ошибка при преобразовании столбца даты: {str(e)}"}
            )
            
        # Проведение анализа с распакованными переменными
        results = analyze_data(df, unpacked_vars)
        
        # Обработка особых случаев возврата
        if isinstance(results, dict) and 'status' in results:
            if results['status'] == 'required_input':
                # Возвращаем список доступных столбцов и сообщение для пользователя
                return JSONResponse(
                    status_code=200,
                    content={
                        "required_input": True,
                        "message": results['message'],
                        "available_columns": results['available_columns']
                    }
                )
            elif results['status'] == 'error':
                # Возвращаем сообщение об ошибке
                return JSONResponse(
                    status_code=400,
                    content={"error": results['message']}
                )
        
        return results
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Ошибка при анализе данных: {str(e)}")
@app.post("/analyze-columns")
async def analyze_columns(file: UploadFile = File(...)):
    """
    Анализ доступных столбцов в загруженном файле
    
    - **file**: Excel или CSV файл с данными
    """
    # Проверка расширения файла
    file_extension = os.path.splitext(file.filename)[1].lower()
    
    try:
        # Чтение данных из файла
        content = await file.read()
        if file_extension == '.csv':
            df = pd.read_csv(BytesIO(content))
        elif file_extension in ['.xlsx', '.xls']:
            df = pd.read_excel(BytesIO(content))
        else:
            raise HTTPException(status_code=400, detail="Поддерживаются только файлы CSV, XLS или XLSX")
        
        # Определение типов столбцов
        column_types = {}
        numeric_columns = []
        
        for column in df.columns:
            if pd.api.types.is_numeric_dtype(df[column]):
                column_types[column] = "numeric"
                numeric_columns.append(column)
            elif pd.api.types.is_datetime64_any_dtype(df[column]):
                column_types[column] = "datetime"
            else:
                try:
                    # Проверяем, можно ли преобразовать в дату
                    pd.to_datetime(df[column])
                    column_types[column] = "potential_datetime"
                except:
                    column_types[column] = "string"
        
        # Возвращаем результаты анализа
        return {
            "columns": df.columns.tolist(),
            "column_types": column_types,
            "numeric_columns": numeric_columns,
            "rows_count": len(df),
            "preview": df.head(5).to_dict(orient="records")
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Ошибка при анализе столбцов: {str(e)}")

@app.get("/api/egrul")
async def get_egrul_data():
    """
    Получение данных ЕГРЮЛ
    """
    return {
        "content": [{
            "dataSved": "2025-02-11",
            "OGRN": "1027700251754",
            "dateOGRN": "2002-09-26",
            "INN": "7710044140",
            "KPP": "770501001",
            "sprOPF": "ОКОНХ",
            "codeOPF": "12300",
            "nameOPF": "Общества с ограниченной ответственностью",
            "fullOrgName": "ОБЩЕСТВО С ОГРАНИЧЕННОЙ ОТВЕТСТВЕННОСТЬЮ \"СИСТЕМА ПБО\"",
            "shortOrgName": "ООО \"СИСТЕМА ПБО\"",
            "region": "77",
            "adres": "115054, Г.МОСКВА , УЛ ВАЛОВАЯ, Д. Д. 26",
            "email": "",
            "statusCode": "",
            "statusName": "",
            "dateClose": "",
            "codeClose": "",
            "causeClose": "",
            "codeNO": "7705",
            "nameNO": "Инспекция Федеральной налоговой службы № 5 по г. Москве",
            "codePF": "",
            "namePF": "",
            "codeFSS": "",
            "nameFSS": "",
            "capitalKind": "УСТАВНЫЙ КАПИТАЛ",
            "capitalAmount": "10694064934.33",
            "ulOGRN": "",
            "ulINN": "",
            "nameUl": "",
            "dirFIO": "ПАРОБЕВ ОЛЕГ МРЬЕВИЧ",
            "dirINN": "774331952837",
            "director": "ГЕНЕРАЛЬНЫЙ ДИРЕКТОР",
            "osnovidKVED": "56.10",
            "nameOKVED": "Деятельность ресторанов и услуги по доставке продуктов питания",
            "verOKVED": "2014",
            "uchreditel": [{
                "OGRN": "1027700251754",
                "uchGRN": "2227705232346",
                "uchINN": "4252003977",
                "uchName": "ОБЩЕСТВО С ОГРАНИЧЕННОЙ ОТВЕТСТВЕННОСТЬЮ \"КЛУБ ОТЕЛЬ\"",
                "uchType": "ULR",
                "neddan": ""
            }]
        }]
    }
@app.get("/health")
async def health_check():
    """
    Проверка работоспособности API
    """
    return {"status": "ok", "version": "1.0.0","Первый запуск Егор":"Оно работает я проверил"}

@app.post("/api/feedback")
async def submit_feedback(feedback: FeedbackForm):
    """
    Обработка формы обратной связи (тестовая версия)
    """
    # Выводим данные в консоль для отладки
    print(f"Получено сообщение от: {feedback.name} <{feedback.email}>")
    print(f"Тема: {feedback.subject}")
    print(f"Сообщение: {feedback.message}")
    
    # Всегда возвращаем успех
    return {"status": "success", "message": "Сообщение успешно отправлено"}

# Генерация тестовых данных для демонстрации (если нет реальных данных)
@app.get("/generate-demo-data")
async def generate_demo_data():
    """
    Генерация тестовых данных для демонстрации API
    """
    try:
        # Создаем тестовые временные ряды
        dates = pd.date_range(start='2013-01-01', end='2023-12-01', freq='MS')
        
        # Симулируем нестационарный ряд (случайное блуждание)
        random_walk = np.random.normal(0, 1, len(dates))
        random_walk = np.cumsum(random_walk)
        
        # Симулируем ряд с трендом
        trend = np.linspace(5, 15, len(dates)) + np.random.normal(0, 0.5, len(dates))
        
        # Создаем датафрейм
        df = pd.DataFrame({
            'Дата': dates,
            'Средневзвешенная_ставка': trend,
            'Средняя_задолженность': random_walk + 50 + np.linspace(0, 30, len(dates)),
            'Ключевая_ставка': np.random.normal(8, 2, len(dates)),
            'Денежные_доходы': np.linspace(25, 50, len(dates)) + np.random.normal(0, 2, len(dates))
        })
        
        # Сохраняем в файл
        with tempfile.NamedTemporaryFile(suffix='.xlsx', delete=False) as temp_file:
            temp_path = temp_file.name
            df.to_excel(temp_path, index=False)
        
        # Отправляем демо-файл пользователю
        return FileResponse(
            path=temp_path,
            filename="demo_data.xlsx",
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Ошибка при генерации демо-данных: {str(e)}")
        
@app.post("/transform-integration-order")
async def transform_integration_order(
    file_id: str = Form(...),
    date_column: str = Form("Дата"),
    transformation_settings: str = Form(...),
    preview: str = Form("false"),
    display_mode: str = Form("all")
):
    """
    Преобразование временных рядов к заданному порядку интеграции с использованием различных методов
    
    - **file_id**: Идентификатор загруженного файла
    - **date_column**: Имя столбца с датами
    - **transformation_settings**: JSON со структурой {переменная: {transformTo: порядок_интеграции, diffMethod: метод_преобразования}}
    - **preview**: Если "true", то возвращает только данные для предпросмотра без сохранения
    - **display_mode**: Режим отображения - "all", "transformed" или "original"
    
    Возможные значения порядка интеграции:
    - "none": Оставить как есть
    - "I(1)": Привести к первому порядку интеграции
    - "I(2)": Привести к второму порядку интеграции
    - "I(3)": Привести к третьему порядку интеграции
    - "I(4)": Привести к четвертому порядку интеграции
    - "normalize": Применить нормализацию
    
    Возможные методы преобразования:
    - "simple": Обычное дифференцирование (разности)
    - "percentage": Процентное изменение
    - "log": Логарифмическое изменение
    - "boxcox": Преобразование Бокса-Кокса
    - "yeojohnson": Преобразование Йео-Джонсона
    - "rank": Ранговая нормализация
    - "auto": Автоматический выбор метода
    - "extreme": Экстремальное преобразование
    """
    # Проверяем наличие файла
    file_path = TEMP_FILES_DIR / file_id
    
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="Файл не найден или срок его хранения истек")
    
    try:
        # Разбираем настройки преобразования
        settings = json.loads(transformation_settings)
        is_preview = preview.lower() == "true"
        
        # Определяем тип файла по расширению
        file_extension = os.path.splitext(str(file_path))[1].lower()
        
        # Чтение данных из файла
        if file_extension == '.csv':
            df = pd.read_csv(file_path)
        elif file_extension in ['.xlsx', '.xls']:
            df = pd.read_excel(file_path)
        else:
            # Пытаемся определить тип файла автоматически
            try:
                df = pd.read_excel(file_path)
            except:
                try:
                    df = pd.read_csv(file_path)
                except:
                    raise HTTPException(status_code=400, detail="Неподдерживаемый формат файла")
        
        # Проверка наличия столбца с датами
        if date_column not in df.columns:
            return {"error": f"Столбец {date_column} не найден в файле", "columns": df.columns.tolist()}
            
        # Преобразование столбца даты
        try:
            df[date_column] = pd.to_datetime(df[date_column])
            # Сортировка по дате для корректности расчетов
            df = df.sort_values(by=date_column)
        except Exception as e:
            return {"error": f"Ошибка при преобразовании столбца даты: {str(e)}"}
        
        # Данные для предпросмотра трансформации
        preview_data = {"transformed_data": {}}
        
        # Создаем новый DataFrame для измененных данных
        transformed_df = df.copy()
        
        # Импортируем функции для преобразования
        from ts_analysis import extreme_transform_to_stationary
        
        # Преобразуем каждую переменную согласно настройкам
        for variable, transform_settings in settings.items():
            # Получаем настройки трансформации
            if isinstance(transform_settings, str):
                # Совместимость со старым форматом
                transform_to = transform_settings
                transform_method = 'simple'
            else:
                # Новый формат с методом преобразования
                transform_to = transform_settings.get('transformTo', 'none')
                transform_method = transform_settings.get('diffMethod', 'simple')
            
            # Пропускаем, если оставляем как есть
            if transform_to == "none":
                continue
                
            # Проверяем, есть ли переменная в DataFrame
            if variable not in df.columns:
                continue
                
            # Получаем временной ряд
            series = df[variable].copy()
            
            # Преобразуем согласно выбранному порядку интеграции
            if transform_to in ["I(1)", "I(2)", "I(3)", "I(4)"]:
                # Получаем порядок интеграции
                order = int(transform_to[2])
                
                # Применяем дифференцирование нужное количество раз
                transformed_series = series
                for i in range(order):
                    if transform_method == 'simple':
                        transformed_series = transformed_series.diff()
                    elif transform_method == 'percentage':
                        transformed_series = transformed_series.pct_change()
                    elif transform_method == 'log':
                        # Логарифмическое дифференцирование
                        if transformed_series.min() <= 0:
                            min_val = transformed_series.min()
                            offset = abs(min_val) + 1
                            transformed_series = np.log(transformed_series + offset)
                        else:
                            transformed_series = np.log(transformed_series)
                        transformed_series = transformed_series.diff()
                
                # Определяем имя для нового столбца
                column_name = f"{variable}_I{order}"
                
            elif transform_to == "normalize":
                # Применяем нормализацию
                transformed_series, transform_info = extreme_transform_to_stationary(
                    series, 
                    method=transform_method, 
                    title=variable
                )
                
                # Определяем имя для нового столбца
                if transform_method == 'auto':
                    # Используем информацию о выбранном методе из transform_info
                    method_suffix = transform_info['transform_method'].split(' ')[0].lower()
                    column_name = f"{variable}_norm_{method_suffix}"
                else:
                    column_name = f"{variable}_norm_{transform_method}"
            
            else:
                # Неизвестный тип преобразования, пропускаем
                continue
            
            # Добавляем преобразованный ряд в DataFrame, только если режим отображения позволяет
            if display_mode in ['all', 'transformed']:
                transformed_df[column_name] = transformed_series
                
                # В режиме "только преобразованные" удаляем исходный ряд, если он не нужен
                if display_mode == 'transformed' and column_name != variable:
                    if variable in transformed_df.columns:
                        transformed_df = transformed_df.drop(columns=[variable])
            
            # Добавляем данные для предпросмотра
            if is_preview:
                # Используем функцию форматирования с учетом режима отображения
                preview_data["transformed_data"][variable] = format_time_series_for_preview(
                    series, 
                    transformed_series, 
                    display_mode=display_mode
                )
        
        # Если режим "только оригинальные", удаляем все преобразованные ряды
        if display_mode == 'original':
            # Находим все столбцы с преобразованиями
            transformed_columns = [col for col in transformed_df.columns if 
                                  col.endswith(('_I1', '_I2', '_I3', '_I4', '_norm')) or
                                  '_norm_' in col]
            if transformed_columns:
                transformed_df = transformed_df.drop(columns=transformed_columns)
        
        # Если это предпросмотр, возвращаем только данные для визуализации
        if is_preview:
            return preview_data
        
        # Сохраняем преобразованные данные в новый файл с правильным расширением
        new_file_id = str(uuid.uuid4())
        new_file_path = TEMP_FILES_DIR / f"{new_file_id}.xlsx"
        
        # Сохраняем файл в том же формате, что и исходный
        if file_extension == '.csv':
            transformed_df.to_csv(new_file_path, index=False)
        else:
            transformed_df.to_excel(new_file_path, index=False)
        
        # Запоминаем время загрузки нового файла
        file_timestamps[new_file_id] = time.time()
        
        # Анализируем столбцы нового файла
        numeric_columns = [col for col in transformed_df.columns if pd.api.types.is_numeric_dtype(transformed_df[col])]
        
        # Готовим данные для обновления localStorage
        updated_columns_data = {
            "columns": transformed_df.columns.tolist(),
            "numeric_columns": numeric_columns,
            "rows_count": len(transformed_df),
            "date_column": date_column,
        }
        
        # Формируем URL для скачивания с сохранением расширения
        download_url = f"http://37.252.23.30:8000/temp-file/{new_file_id}.xlsx"
        
        # Возвращаем результаты
        return {
            "success": True,
            "message": "Преобразование успешно выполнено",
            "updated_file_id": f"{new_file_id}",  # без расширения
            "updated_columns_data": updated_columns_data,
            "download_url": download_url,
            "display_mode": display_mode
        }
            
    except Exception as e:
        import traceback
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Ошибка при преобразовании данных: {str(e)}")
        
@app.post("/build-varx-model")
async def build_varx_model_endpoint(
    file_id: str = Form(...),
    date_column: str = Form("Дата"),
    endogenous_vars: Optional[Union[List[str], str]] = Form(None),
    exogenous_vars: Optional[Union[List[str], str]] = Form(None),
    lags: int = Form(1),
    train_size: float = Form(0.8),
    forecast_periods: int = Form(6),
    forecast_unit: str = Form("months")
):
    """
    Построение VARX модели на основе загруженных данных с прогнозированием в будущее
    
    - **file_id**: Идентификатор файла, полученный при загрузке
    - **date_column**: Имя столбца с датами (по умолчанию "Дата")
    - **endogenous_vars**: Список имен эндогенных переменных для модели
    - **exogenous_vars**: Список имен экзогенных переменных для модели
    - **lags**: Количество лагов для модели (по умолчанию 1)
    - **train_size**: Доля данных для обучения модели (по умолчанию 0.8)
    - **forecast_periods**: Количество периодов для прогнозирования в будущее (по умолчанию 6)
    - **forecast_unit**: Единица времени для прогнозирования: months, quarters, years (по умолчанию months)
    """
    # Распаковываем переменные из возможных форматов
    unpacked_endog = unpack_variables(endogenous_vars)
    unpacked_exog = unpack_variables(exogenous_vars)
    
    print(f"VARX model - file_id: {file_id}")
    print(f"VARX model - date_column: {date_column}")
    print(f"VARX model - endogenous_vars: {unpacked_endog}")
    print(f"VARX model - exogenous_vars: {unpacked_exog}")
    print(f"VARX model - lags: {lags}")
    print(f"VARX model - train_size: {train_size}")
    print(f"VARX model - forecast_periods: {forecast_periods}")
    print(f"VARX model - forecast_unit: {forecast_unit}")
    
    # Проверяем наличие хотя бы одной эндогенной переменной
    if not unpacked_endog:
        return JSONResponse(
            status_code=400,
            content={"error": "Необходимо указать хотя бы одну эндогенную переменную"}
        )
    
    file_path = TEMP_FILES_DIR / file_id
    
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="Файл не найден или срок его хранения истек")
    
    try:
        # Определяем тип файла по расширению
        file_extension = os.path.splitext(str(file_path))[1].lower()
        
        # Чтение данных из файла
        if file_extension == '.csv':
            df = pd.read_csv(file_path)
        elif file_extension in ['.xlsx', '.xls']:
            df = pd.read_excel(file_path)
        else:
            # Пытаемся определить тип файла автоматически
            try:
                df = pd.read_excel(file_path)
            except:
                try:
                    df = pd.read_csv(file_path)
                except:
                    raise HTTPException(status_code=400, detail="Неподдерживаемый формат файла")
        
        # Проверка наличия столбца с датами
        if date_column not in df.columns:
            return JSONResponse(
                status_code=400,
                content={"error": f"Столбец {date_column} не найден в файле", "columns": df.columns.tolist()}
            )
            
        # Преобразование столбца даты
        try:
            df[date_column] = pd.to_datetime(df[date_column])
            df.set_index(date_column, inplace=True)
        except Exception as e:
            return JSONResponse(
                status_code=400,
                content={"error": f"Ошибка при преобразовании столбца даты: {str(e)}"}
            )
        
        # Проверяем, что все указанные переменные есть в данных
        all_vars = unpacked_endog + unpacked_exog if unpacked_exog else unpacked_endog
        missing_vars = [var for var in all_vars if var not in df.columns]
        if missing_vars:
            return JSONResponse(
                status_code=400,
                content={
                    "error": f"Следующие переменные не найдены в данных: {', '.join(missing_vars)}",
                    "available_columns": df.columns.tolist()
                }
            )
        
        # Строим VARX/ARX модель с прогнозированием в будущее
        from ts_analysis import build_varx_model_with_future_forecast
        
        model_results = build_varx_model_with_future_forecast(
            df, 
            unpacked_endog, 
            unpacked_exog, 
            lags=lags,
            train_size=train_size,
            forecast_periods=forecast_periods,
            forecast_unit=forecast_unit
        )
        
        return model_results
        
    except Exception as e:
        import traceback
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Ошибка при построении модели: {str(e)}")

@app.get("/download-varx-report/{file_id}")
async def download_varx_report(
    file_id: str,
    endogenous: str,
    exogenous: str = "",
    lags: int = 1,
    forecast_periods: int = 6,
    forecast_unit: str = "months"
):
    """
    Скачивание отчета по VARX/ARX модели с прогнозированием в будущее
    
    - **file_id**: Идентификатор файла с данными
    - **endogenous**: JSON-строка с именами эндогенных переменных
    - **exogenous**: JSON-строка с именами экзогенных переменных
    - **lags**: Количество лагов модели
    - **forecast_periods**: Количество периодов для прогнозирования в будущее
    - **forecast_unit**: Единица времени для прогнозирования
    """
    try:
        # Распаковываем переменные
        endogenous_vars = json.loads(endogenous)
        exogenous_vars = json.loads(exogenous) if exogenous else []
        
        file_path = TEMP_FILES_DIR / file_id
        
        if not file_path.exists():
            raise HTTPException(status_code=404, detail="Файл не найден или срок его хранения истек")
        
        # Загружаем данные
        if os.path.splitext(str(file_path))[1].lower() == '.csv':
            df = pd.read_csv(file_path)
        else:
            df = pd.read_excel(file_path)
        
        # Преобразуем столбец даты
        if 'Дата' in df.columns:
            df['Дата'] = pd.to_datetime(df['Дата'])
            df.set_index('Дата', inplace=True)
        
        # Строим модель с прогнозированием в будущее
        from ts_analysis import build_varx_model_with_future_forecast
        model_results = build_varx_model_with_future_forecast(
            df, 
            endogenous_vars, 
            exogenous_vars, 
            lags=lags,
            forecast_periods=forecast_periods,
            forecast_unit=forecast_unit
        )
        
        # Создаем отчет в Word
        from docx import Document
        from docx.shared import Inches
        import matplotlib.pyplot as plt
        import base64
        import io
        
        doc = Document()
        
        # Определяем тип модели
        model_type = "ARX" if len(endogenous_vars) == 1 else "VARX"
        
        # Добавляем заголовок
        doc.add_heading(f'Отчет по {model_type} модели с прогнозированием в будущее', 0)
        
        # Добавляем текущую дату
        current_date = datetime.now().strftime("%d.%m.%Y")
        doc.add_paragraph(f'Дата создания отчета: {current_date}')
        
        # Информация о модели
        doc.add_heading('Информация о модели', 1)
        
        # Таблица с основными параметрами модели
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Параметр'
        hdr_cells[1].text = 'Значение'
        
        # Заполняем таблицу
        params = [
            ('Тип модели', model_type),
            ('Эндогенные переменные', ', '.join(endogenous_vars)),
            ('Экзогенные переменные', ', '.join(exogenous_vars) if exogenous_vars else 'Не используются'),
            ('Количество лагов', str(lags))
        ]
        
        # Добавляем информацию о прогнозном периоде
        unit_names = {
            'months': 'месяцев',
            'quarters': 'кварталов', 
            'years': 'лет'
        }
        unit_name = unit_names.get(forecast_unit, forecast_unit)
        params.append(('Период прогнозирования', f'{forecast_periods} {unit_name}'))
        
        # Добавляем параметры, если они есть в результатах
        if 'model_info' in model_results:
            model_info = model_results['model_info']
            
            if 'aic' in model_info and model_info['aic'] is not None:
                params.append(('AIC', f"{model_info['aic']:.4f}"))
            
            if 'bic' in model_info and model_info['bic'] is not None:
                params.append(('BIC', f"{model_info['bic']:.4f}"))
            
            if 'hqic' in model_info and model_info['hqic'] is not None:
                params.append(('HQIC', f"{model_info['hqic']:.4f}"))
            
            if 'train_size' in model_info:
                params.append(('Размер обучающей выборки', f"{model_info['train_size']} наблюдений"))
            
            if 'test_size' in model_info:
                params.append(('Размер тестовой выборки', f"{model_info['test_size']} наблюдений"))
        
        for param, value in params:
            row_cells = table.add_row().cells
            row_cells[0].text = param
            row_cells[1].text = value
        
        # Добавляем раздел о выборе оптимального числа лагов, если есть
        if 'diagnostics' in model_results and 'lag_analysis' in model_results['diagnostics']:
            doc.add_heading('Анализ оптимального числа лагов', 1)
            
            lag_analysis = model_results['diagnostics']['lag_analysis']
            
            # Создаем таблицу с оптимальными лагами
            lag_table = doc.add_table(rows=1, cols=2)
            lag_table.style = 'Table Grid'
            
            hdr_cells = lag_table.rows[0].cells
            hdr_cells[0].text = 'Переменная'
            hdr_cells[1].text = 'Оптимальное число лагов'
            
            # Заполняем таблицу для каждой переменной
            for var, lag in lag_analysis['optimal_lags'].items():
                row_cells = lag_table.add_row().cells
                row_cells[0].text = var
                row_cells[1].text = str(lag)
            
            # Добавляем рекомендации
            doc.add_paragraph(f"Рекомендуемое количество лагов (максимальное из оптимальных): {lag_analysis['optimal_lag_final']}")
            doc.add_paragraph(f"В модели использовано: {lag_analysis['chosen_lag']} лагов")
        
        # Добавляем графики ACF и PACF, если они есть
        if 'plots' in model_results and 'acf_pacf_plots' in model_results['plots']:
            doc.add_heading('Графики автокорреляционных функций', 1)
            
            for i, plot_base64 in enumerate(model_results['plots']['acf_pacf_plots']):
                # Преобразуем base64 в изображение
                image_data = base64.b64decode(plot_base64)
                image_stream = io.BytesIO(image_data)
                
                # Добавляем изображение в документ
                try:
                    doc.add_picture(image_stream, width=Inches(6))
                    if i < len(endogenous_vars):
                        doc.add_paragraph(f'ACF и PACF для переменной "{endogenous_vars[i]}"')
                    else:
                        doc.add_paragraph(f'ACF и PACF график {i+1}')
                except Exception as e:
                    doc.add_paragraph(f'Не удалось добавить изображение графика: {str(e)}')
        
        # Добавляем метрики прогнозирования на валидационных данных, если они есть
        if 'validation' in model_results and 'metrics' in model_results['validation']:
            doc.add_heading('Метрики качества прогноза (на тестовых данных)', 1)
            
            metrics_table = doc.add_table(rows=1, cols=5)
            metrics_table.style = 'Table Grid'
            
            hdr_cells = metrics_table.rows[0].cells
            hdr_cells[0].text = 'Переменная'
            hdr_cells[1].text = 'MSE'
            hdr_cells[2].text = 'MAE'
            hdr_cells[3].text = 'RMSE'
            hdr_cells[4].text = 'MAPE'
            
            metrics = model_results['validation']['metrics']
            
            for variable in endogenous_vars:
                if variable in metrics['mse']:
                    row_cells = metrics_table.add_row().cells
                    row_cells[0].text = variable
                    row_cells[1].text = f"{metrics['mse'][variable]:.4f}"
                    row_cells[2].text = f"{metrics['mae'][variable]:.4f}"
                    
                    # Добавляем RMSE, если есть
                    if 'rmse' in metrics and variable in metrics['rmse']:
                        row_cells[3].text = f"{metrics['rmse'][variable]:.4f}"
                    else:
                        row_cells[3].text = 'N/A'
                    
                    # Добавляем MAPE, если есть
                    if 'mape' in metrics and variable in metrics['mape'] and not np.isnan(metrics['mape'][variable]):
                        row_cells[4].text = f"{metrics['mape'][variable]:.2f}%"
                    else:
                        row_cells[4].text = 'N/A'
        
        # Добавляем раздел с прогнозными значениями
        if 'forecasts' in model_results and 'future_forecast' in model_results['forecasts']:
            doc.add_heading('Прогнозные значения в будущее', 1)
            
            future_data = model_results['forecasts']['future_forecast']
            
            # Создаем таблицу с прогнозными значениями
            forecast_table = doc.add_table(rows=1, cols=len(endogenous_vars) + 1)
            forecast_table.style = 'Table Grid'
            
            # Заголовки таблицы
            hdr_cells = forecast_table.rows[0].cells
            hdr_cells[0].text = 'Дата'
            for i, var in enumerate(endogenous_vars):
                hdr_cells[i + 1].text = var
            
            # Заполняем таблицу прогнозными значениями
            for i, date in enumerate(future_data['dates']):
                row_cells = forecast_table.add_row().cells
                row_cells[0].text = date
                for j, var in enumerate(endogenous_vars):
                    if var in future_data['values'] and i < len(future_data['values'][var]):
                        row_cells[j + 1].text = f"{future_data['values'][var][i]:.4f}"
                    else:
                        row_cells[j + 1].text = 'N/A'
        
        # Добавляем графики прогнозов, если они есть
        if 'plots' in model_results and 'forecast_plots' in model_results['plots']:
            doc.add_heading('Графики прогнозов с будущими значениями', 1)
            
            plots = model_results['plots']['forecast_plots']
            
            for i, plot_base64 in enumerate(plots):
                # Преобразуем base64 в изображение
                image_data = base64.b64decode(plot_base64)
                image_stream = io.BytesIO(image_data)
                
                # Добавляем изображение в документ
                try:
                    doc.add_picture(image_stream, width=Inches(6))
                    if i < len(endogenous_vars):
                        doc.add_paragraph(f'Прогноз для переменной "{endogenous_vars[i]}" с будущими значениями')
                    else:
                        doc.add_paragraph(f'Прогноз {i+1} с будущими значениями')
                except Exception as e:
                    doc.add_paragraph(f'Не удалось добавить изображение графика: {str(e)}')
        
        # Сохраняем документ во временный файл
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
            temp_path = temp_file.name
            doc.save(temp_path)
        
        # Возвращаем файл для скачивания
        return FileResponse(
            path=temp_path,
            filename=f"{model_type}_forecast_report_{current_date.replace('.', '-')}.docx",
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
    except Exception as e:
        import traceback
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Ошибка при создании отчета: {str(e)}")
        
# Добавьте следующий класс для описания структуры запроса
class ReportRequest(BaseModel):
    file_id: str
    format: str = "docx"
    tests: Dict[str, Union[bool, Dict[str, Any]]] = {}
    models: Dict[str, Dict[str, Any]] = {}

@app.post("/generate-report")
async def generate_report(request: Request):
    """
    Генерация отчета на основе выбранных тестов и моделей
    
    Поддерживаемые форматы: DOCX
    """
    try:
        # Получаем JSON-данные из тела запроса
        data = await request.json()
        
        # Извлекаем параметры
        file_id = data.get('file_id')
        report_format = data.get('format', 'docx')
        tests = data.get('tests', {})
        models = data.get('models', {})
        
        # Проверка наличия файла
        file_path = TEMP_FILES_DIR / file_id
        if not file_path.exists():
            raise HTTPException(status_code=404, detail="Файл не найден или срок его хранения истек")
        
        # Определяем тип файла по расширению
        file_extension = os.path.splitext(str(file_path))[1].lower()
        
        # Чтение данных из файла
        try:
            if file_extension == '.csv':
                df = pd.read_csv(file_path)
            elif file_extension in ['.xlsx', '.xls']:
                df = pd.read_excel(file_path)
            else:
                # Пытаемся определить тип файла автоматически
                try:
                    df = pd.read_excel(file_path)
                except:
                    try:
                        df = pd.read_csv(file_path)
                    except:
                        raise HTTPException(status_code=400, detail="Неподдерживаемый формат файла")
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Ошибка при чтении файла: {str(e)}")
        
        # Подготовка данных для анализа
        try:
            date_column = "Дата"  # По умолчанию используем Дата
            
            # Проверка наличия столбца с датами
            if date_column not in df.columns:
                # Пытаемся найти столбец с датами
                date_columns = [col for col in df.columns if 'дата' in col.lower() or 'date' in col.lower()]
                if date_columns:
                    date_column = date_columns[0]
                else:
                    raise HTTPException(status_code=400, content={"error": f"Столбец с датами не найден"})
            
            # Преобразование столбца даты
            df[date_column] = pd.to_datetime(df[date_column])
            df.set_index(date_column, inplace=True)
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Ошибка при подготовке данных: {str(e)}")
        
        # Проверка и обработка выбранных моделей
        if 'varx' in models:
            varx_params = models['varx']
            
            # Если параметры VARX указаны как булево значение, преобразуем в словарь
            if isinstance(varx_params, bool):
                varx_params = {}
                models['varx'] = varx_params
            
            # Если не указаны эндогенные и экзогенные переменные, добавляем их в параметры
            if 'endogenous_vars' not in varx_params or not varx_params['endogenous_vars']:
                # Получаем числовые колонки
                numeric_columns = [col for col in df.columns if pd.api.types.is_numeric_dtype(df[col])]
                
                # Берем первые две колонки как эндогенные (или одну, если только одна доступна)
                if len(numeric_columns) > 0:
                    varx_params['endogenous_vars'] = numeric_columns[:min(2, len(numeric_columns))]
                else:
                    raise HTTPException(status_code=400, detail="Не найдены числовые колонки для эндогенных переменных")
            
            # Если экзогенные переменные не указаны, оставляем пустой список
            if 'exogenous_vars' not in varx_params:
                varx_params['exogenous_vars'] = []
        
        # Создание отчета в зависимости от выбранных тестов и моделей
        if report_format == 'docx':
            # Используем функцию create_comprehensive_report из ts_analysis
            from ts_analysis import create_comprehensive_report
            
            # Подготовка параметров для отчета
            report_params = {
                'date_column': date_column,
                'tests': tests,
                'models': models
            }
            
            # Создаем отчет
            report_path = create_comprehensive_report(df, report_params)
            
            # Отправляем файл в ответе
            return FileResponse(
                path=report_path,
                filename=f"Аналитический_отчет.docx",
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        else:
            raise HTTPException(status_code=400, detail=f"Формат {report_format} не поддерживается")
    
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Ошибка при создании отчета: {str(e)}")
# Установка URL API для формирования URL скачивания
API_URL = "http://37.252.23.30:8000" 
if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
